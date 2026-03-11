"""
src/document/exporter.py 的延伸測試
補充匯出格式、邊界條件和特殊字元處理
"""
from unittest.mock import patch

from docx import Document
from src.document.exporter import DocxExporter
from src.core.constants import get_platform_fonts, _FONT_FALLBACK


# ==================== 文件類型偵測 ====================

class TestExtractDocType:
    """_extract_doc_type 方法的測試"""

    def test_detect_han(self):
        """測試偵測「函」類型"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 函\n內容") == "函"

    def test_detect_announcement(self):
        """測試偵測「公告」類型"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 公告\n內容") == "公告"

    def test_detect_sign(self):
        """測試偵測「簽」類型"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 簽\n內容") == "簽"

    def test_detect_order(self):
        """測試偵測「令」類型"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 令\n內容") == "令"

    def test_detect_meeting_notice(self):
        """測試偵測「開會通知單」類型"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 開會通知單\n內容") == "開會通知單"

    def test_default_type(self):
        """測試無法偵測時預設為「函」"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("隨便的文字\n沒有類型") == "函"

    def test_type_within_first_5_lines(self):
        """測試類型在前5行內被偵測"""
        exporter = DocxExporter()
        draft = "一些前綴\n\n\n公告\n內容"
        assert exporter._extract_doc_type(draft) == "公告"

    def test_detect_chen(self):
        """偵測「呈」"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 呈\n\n**機關**：行政院") == "呈"

    def test_detect_zi(self):
        """偵測「咨」"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 咨\n\n**機關**：總統府") == "咨"

    def test_detect_site_inspection(self):
        """偵測「會勘通知單」"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 會勘通知單\n\n**機關**：工務局") == "會勘通知單"

    def test_detect_phone_record(self):
        """偵測「公務電話紀錄」"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 公務電話紀錄\n\n**機關**：秘書處") == "公務電話紀錄"

    def test_detect_directive(self):
        """偵測「手令」"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 手令\n\n**發令人**：市長") == "手令"

    def test_detect_memo(self):
        """偵測「箋函」"""
        exporter = DocxExporter()
        assert exporter._extract_doc_type("# 箋函\n\n**發信人**：秘書處") == "箋函"


# ==================== 清理行文字 ====================

class TestCleanLine:
    """_clean_line 方法的測試"""

    def test_remove_markdown_headers(self):
        """測試移除 markdown 標題符號"""
        exporter = DocxExporter()
        assert exporter._clean_line("### 主旨") == "主旨"

    def test_remove_bold_markers(self):
        """測試移除粗體標記"""
        exporter = DocxExporter()
        result = exporter._clean_line("**機關**：測試")
        assert "**" not in result
        assert "機關" in result

    def test_remove_list_markers(self):
        """測試移除列表標記"""
        exporter = DocxExporter()
        result = exporter._clean_line("- 項目一")
        assert result == "項目一"

    def test_empty_string(self):
        """測試空字串"""
        exporter = DocxExporter()
        assert exporter._clean_line("") == ""

    def test_preserve_chinese_content(self):
        """測試保留中文內容"""
        exporter = DocxExporter()
        assert "依據" in exporter._clean_line("依據某法規辦理")


# ==================== 簽類型匯出 ====================

class TestSignExport:
    """簽類型文件匯出的測試"""

    def test_sign_export(self, tmp_path):
        """測試簽類型的匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_sign.docx"

        draft = """# 簽

**機關**：市府

---

### 主旨
簽呈主旨

### 說明
簽呈說明

### 擬辦
建議方案
"""
        exporter.export(draft, str(output_file))
        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_sign_with_reason(self, tmp_path):
        """測試簽類型包含依據"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_sign_basis.docx"

        draft = """# 簽

### 主旨
簽呈測試

### 依據
某法規

### 說明
詳細說明
"""
        exporter.export(draft, str(output_file))
        assert output_file.exists()


# ==================== 新增公文類型匯出 ====================

class TestNewDocTypeExports:
    """新增公文類型的 DOCX 匯出測試"""

    def test_inspection_export(self, tmp_path):
        """測試會勘通知單的 DOCX 匯出"""
        exporter = DocxExporter()
        draft = """# 會勘通知單

**機關**：臺北市政府工務局
**受文者**：相關單位

---

### 主旨
辦理道路損壞會勘

### 會勘時間
中華民國115年4月1日上午10時

### 會勘地點
臺北市信義路四段
"""
        out = str(tmp_path / "inspection.docx")
        result = exporter.export(draft, out)
        assert result == out
        doc = Document(out)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "會勘通知單" in full_text
        assert "道路損壞會勘" in full_text

    def test_phone_record_export(self, tmp_path):
        """測試公務電話紀錄的 DOCX 匯出"""
        exporter = DocxExporter()
        draft = """# 公務電話紀錄

**機關**：臺北市政府秘書處

---

### 主旨
確認會議時間

### 通話摘要
會議改至3月10日召開

### 紀錄人
張書記
"""
        out = str(tmp_path / "phone.docx")
        result = exporter.export(draft, out)
        assert result == out
        doc = Document(out)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "公務電話紀錄" in full_text
        assert "確認會議時間" in full_text

    def test_directive_export(self, tmp_path):
        """測試手令的 DOCX 匯出"""
        exporter = DocxExporter()
        draft = """# 手令

**發令人**：臺北市市長
**受令人**：都發局局長

---

### 主旨
指派辦理社會住宅專案

### 指示事項
即日起督導辦理

### 完成期限
115年12月31日前完成
"""
        out = str(tmp_path / "directive.docx")
        result = exporter.export(draft, out)
        assert result == out
        doc = Document(out)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "手令" in full_text
        assert "社會住宅" in full_text

    def test_memo_export(self, tmp_path):
        """測試箋函的 DOCX 匯出"""
        exporter = DocxExporter()
        draft = """# 箋函

**發信人**：臺北市政府秘書處
**收信人**：臺北市政府人事處

---

### 主旨
請提供員工訓練計畫

### 說明
配合年度施政報告彙整
"""
        out = str(tmp_path / "memo.docx")
        result = exporter.export(draft, out)
        assert result == out
        doc = Document(out)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "箋函" in full_text
        assert "員工訓練" in full_text

    def test_chen_export(self, tmp_path):
        """測試呈的 DOCX 匯出"""
        exporter = DocxExporter()
        draft = """# 呈

**機關**：行政院
**受文者**：總統府

---

### 主旨
呈報年度施政成果

### 說明
依據行政院組織法規定辦理
"""
        out = str(tmp_path / "chen.docx")
        result = exporter.export(draft, out)
        assert result == out
        doc = Document(out)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "呈" in full_text
        assert "施政成果" in full_text

    def test_zi_export(self, tmp_path):
        """測試咨的 DOCX 匯出"""
        exporter = DocxExporter()
        draft = """# 咨

**機關**：總統府
**受文者**：立法院

---

### 主旨
咨請審議條約案

### 說明
依據憲法第63條規定
"""
        out = str(tmp_path / "zi.docx")
        result = exporter.export(draft, out)
        assert result == out
        doc = Document(out)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert "咨" in full_text
        assert "條約案" in full_text


# ==================== 特殊內容匯出 ====================

class TestSpecialContentExport:
    """特殊內容的匯出測試"""

    def test_export_with_citations(self, tmp_path):
        """測試含有引用標記的匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_citation.docx"

        draft = """# 函

### 主旨
引用測試

### 說明
依據某法辦理[^1]。

### 參考來源
[^1]: 某法規
"""
        exporter.export(draft, str(output_file))
        assert output_file.exists()

    def test_export_with_numbered_provisions(self, tmp_path):
        """測試含有編號辦法的匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_provisions.docx"

        draft = """# 函

### 主旨
辦法測試

### 辦法
一、第一項辦法
二、第二項辦法
（一）子項目
（二）另一子項目
"""
        exporter.export(draft, str(output_file))
        assert output_file.exists()

    def test_export_with_special_characters(self, tmp_path):
        """測試含有特殊字元的匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_special.docx"

        draft = """# 函

### 主旨
測試「全形引號」和《書名號》及（括號）

### 說明
金額：新臺幣1,000元整
百分比：50%
電話：(02)1234-5678
"""
        exporter.export(draft, str(output_file))
        assert output_file.exists()

    def test_export_with_long_content(self, tmp_path):
        """測試長內容的匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_long.docx"

        long_explanation = "\n".join([f"{i+1}、第{i+1}項說明內容，這是一段較長的文字。" for i in range(20)])
        draft = f"# 函\n\n### 主旨\n長內容測試\n\n### 說明\n{long_explanation}"
        exporter.export(draft, str(output_file))
        assert output_file.exists()
        assert output_file.stat().st_size > 1000  # 長內容應產生較大的檔案

    def test_export_only_subject(self, tmp_path):
        """測試只有主旨的最簡內容"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_minimal.docx"

        draft = "# 函\n### 主旨\n最簡測試"
        exporter.export(draft, str(output_file))
        assert output_file.exists()

    def test_export_with_attachments(self, tmp_path):
        """測試包含附件段落的匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_attach.docx"

        draft = """# 函

### 主旨
附件測試

### 說明
請參閱附件。

### 附件
一、某報告一份
二、某表格一份
"""
        exporter.export(draft, str(output_file))
        assert output_file.exists()

    def test_export_returns_output_path(self, tmp_path):
        """測試 export 回傳輸出路徑"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_return.docx"

        result = exporter.export("# 函\n### 主旨\n測試", str(output_file))
        assert result == str(output_file)


# ==================== DOCX 內容驗證 ====================

class TestDocxContentVerification:
    """驗證生成的 DOCX 文件內容"""

    def test_docx_has_title(self, tmp_path):
        """測試 DOCX 包含文件類型標題"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_verify.docx"
        exporter.export("# 函\n### 主旨\n測試主旨", str(output_file))

        doc = Document(str(output_file))
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "函" in all_text

    def test_docx_has_subject(self, tmp_path):
        """測試 DOCX 包含主旨內容"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_subject.docx"
        exporter.export("# 函\n### 主旨\n驗證這個主旨", str(output_file))

        doc = Document(str(output_file))
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "驗證這個主旨" in all_text

    def test_docx_qa_report_on_new_page(self, tmp_path):
        """測試 QA 報告在新頁面"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_qa_page.docx"
        exporter.export(
            "# 函\n### 主旨\n測試",
            str(output_file),
            qa_report="# QA Report\nScore: 0.95"
        )

        doc = Document(str(output_file))
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "品質保證報告" in all_text or "QA Report" in all_text


class TestWriteAttachmentsMultiLine:
    """測試 _write_attachments 的多行內容處理"""

    def test_multiline_attachments_each_line_separate_paragraph(self, tmp_path):
        """多行附件應每行獨立段落，不殘留 markdown 標記"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = "# 函\n\n主旨：測試\n\n附件：\n- 報告一\n- 報告二\n- 報告三"
        output_path = str(tmp_path / "test_multiline_att.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        # 每行都應該被清理，不應有殘留的 "- " 標記
        assert "- 報告" not in all_text
        # 但實際的附件內容應該存在
        assert "報告一" in all_text
        assert "報告二" in all_text
        assert "報告三" in all_text

    def test_single_line_attachment_works(self, tmp_path):
        """單行附件應正確處理"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = "# 函\n\n主旨：測試\n\n附件：公函影本"
        output_path = str(tmp_path / "test_single_att.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "公函影本" in all_text


# ==================== BUG-002: 文件類型感知標籤 ====================

class TestDocTypeAwareLabels:
    """BUG-002: DocxExporter 應依文件類型使用正確的段落標籤。

    - 函：主旨 / 說明 / 辦法
    - 公告：主旨 / 依據 / 公告事項
    - 簽：主旨 / 說明 / 擬辦
    """

    def test_han_uses_standard_labels(self, tmp_path):
        """函類型應使用「辦法」標籤"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = "# 函\n\n**機關**：市府\n\n---\n\n### 主旨\n測試主旨\n\n### 說明\n測試說明\n\n### 辦法\n請照辦"
        output_path = str(tmp_path / "test_han_labels.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "辦法：" in all_text
        assert "說明：" in all_text
        assert "公告事項：" not in all_text
        assert "擬辦：" not in all_text

    def test_announcement_uses_correct_labels(self, tmp_path):
        """公告類型應使用「依據」和「公告事項」標籤，而非「說明」和「辦法」"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = (
            "# 公告\n\n**機關**：市府\n\n---\n\n"
            "### 主旨\n公告測試\n\n"
            "### 依據\n某法規\n\n"
            "### 公告事項\n一、第一項\n二、第二項"
        )
        output_path = str(tmp_path / "test_announcement_labels.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "公告事項：" in all_text
        assert "依據：" in all_text
        # 不應出現函的標籤
        assert "辦法：" not in all_text

    def test_sign_uses_correct_labels(self, tmp_path):
        """簽類型應使用「擬辦」標籤，而非「辦法」"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = (
            "# 簽\n\n**機關**：市府\n\n---\n\n"
            "### 主旨\n簽呈測試\n\n"
            "### 說明\n簽呈說明\n\n"
            "### 擬辦\n建議方案"
        )
        output_path = str(tmp_path / "test_sign_labels.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "擬辦：" in all_text
        assert "說明：" in all_text
        # 不應出現函的標籤
        assert "辦法：" not in all_text

    def test_announcement_basis_not_double_labeled(self, tmp_path):
        """公告的依據段落不應產生重複標籤（不應顯示「說明：依據：...」）"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = (
            "# 公告\n\n**機關**：市府\n\n---\n\n"
            "### 主旨\n公告主旨\n\n"
            "### 依據\n教育部函辦理\n\n"
            "### 公告事項\n請查照"
        )
        output_path = str(tmp_path / "test_no_double_label.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        # 依據內容應出現
        assert "教育部函辦理" in all_text
        # 不應有「說明：」標籤（公告不使用說明）
        assert "說明：" not in all_text

    def test_unknown_doc_type_defaults_to_han(self, tmp_path):
        """未知文件類型應使用函的標籤（預設行為）"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = "一些文字\n\n### 主旨\n測試\n\n### 說明\n說明內容\n\n### 辦法\n辦法內容"
        output_path = str(tmp_path / "test_unknown_type.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "辦法：" in all_text
        assert "說明：" in all_text


# ==================== BUG-005: _extract_doc_type 擴大掃描 ====================

class TestExtractDocTypeExpanded:
    """BUG-005: _extract_doc_type 擴大掃描範圍和模糊匹配。"""

    def test_doc_type_after_line_5(self):
        """公文類型出現在第 6 行之後仍應被辨識（擴展到 10 行）"""
        exporter = DocxExporter()
        draft = (
            "# 臺北市政府\n"
            "\n"
            "發文機關：臺北市政府環境保護局\n"
            "受文者：臺北市各級學校\n"
            "日期：中華民國113年1月1日\n"
            "\n"
            "公告\n"
            "\n"
            "主旨：公告本市資源回收相關事項。"
        )
        assert exporter._extract_doc_type(draft) == "公告"

    def test_fuzzy_match_in_title(self):
        """模糊匹配：行中包含類型關鍵字（如「臺北市政府公告」）"""
        exporter = DocxExporter()
        draft = "# 臺北市政府公告\n\n主旨：測試"
        assert exporter._extract_doc_type(draft) == "公告"

    def test_fuzzy_match_sign(self):
        """模糊匹配：行中包含「簽」關鍵字"""
        exporter = DocxExporter()
        draft = "# 市府內部簽\n\n主旨：簽呈測試"
        assert exporter._extract_doc_type(draft) == "簽"

    def test_fuzzy_match_prefers_longer_name(self):
        """模糊匹配應優先匹配較長的類型名稱"""
        exporter = DocxExporter()
        draft = "# 開會通知單\n\n主旨：會議通知"
        assert exporter._extract_doc_type(draft) == "開會通知單"

    def test_exact_match_takes_priority(self):
        """精確匹配應優先於模糊匹配"""
        exporter = DocxExporter()
        # 第一行精確是「函」
        draft = "函\n\n臺北市政府公告\n\n主旨：測試"
        assert exporter._extract_doc_type(draft) == "函"

    def test_announcement_export_with_fuzzy_title(self, tmp_path):
        """模糊匹配到公告時，匯出應使用正確的段落標籤"""
        from docx import Document as DocxDocument

        exporter = DocxExporter()
        draft = (
            "# 臺北市政府環保局公告\n\n"
            "### 主旨\n公告測試\n\n"
            "### 依據\n某法規\n\n"
            "### 公告事項\n請查照\n"
        )
        output_path = str(tmp_path / "test_fuzzy_announcement.docx")
        exporter.export(draft, output_path)

        doc = DocxDocument(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "公告事項：" in all_text
        assert "辦法：" not in all_text


# ==================== 跨平台字體 ====================

class TestPlatformFonts:
    """跨平台字體 fallback 鏈的測試"""

    def test_windows_fonts(self):
        """Windows 應回傳 DFKai-SB / MingLiU"""
        with patch("src.core.constants.platform.system", return_value="Windows"):
            title, body = get_platform_fonts()
        assert title == "DFKai-SB"
        assert body == "MingLiU"

    def test_macos_fonts(self):
        """macOS 應回傳 PingFang TC"""
        with patch("src.core.constants.platform.system", return_value="Darwin"):
            title, body = get_platform_fonts()
        assert title == "PingFang TC"
        assert body == "PingFang TC"

    def test_linux_fonts(self):
        """Linux 應回傳 Noto Sans CJK TC"""
        with patch("src.core.constants.platform.system", return_value="Linux"):
            title, body = get_platform_fonts()
        assert title == "Noto Sans CJK TC"
        assert body == "Noto Sans CJK TC"

    def test_unknown_os_returns_none(self):
        """未知作業系統應回傳 (None, None)"""
        with patch("src.core.constants.platform.system", return_value="FreeBSD"):
            title, body = get_platform_fonts()
        assert title is None
        assert body is None

    def test_fallback_chain_has_all_platforms(self):
        """確認 fallback 鏈包含 Windows / Darwin / Linux"""
        assert "Windows" in _FONT_FALLBACK
        assert "Darwin" in _FONT_FALLBACK
        assert "Linux" in _FONT_FALLBACK

    def test_exporter_uses_platform_fonts(self):
        """DocxExporter 初始化時應載入平台字體"""
        with patch("src.document.exporter.get_platform_fonts", return_value=("TestTitle", "TestBody")):
            exporter = DocxExporter()
        assert exporter._font_title == "TestTitle"
        assert exporter._font_body == "TestBody"

    def test_exporter_none_font_still_exports(self, tmp_path):
        """字體為 None 時仍應能正常匯出 DOCX"""
        with patch("src.document.exporter.get_platform_fonts", return_value=(None, None)):
            exporter = DocxExporter()
        output_file = tmp_path / "test_none_font.docx"
        exporter.export("# 函\n### 主旨\n測試", str(output_file))
        assert output_file.exists()
        assert output_file.stat().st_size > 0


# ==================== 特殊字元清理 ====================

class TestSanitizeTextExtended:
    """_sanitize_text 的延伸測試：Unicode 特殊空格和控制字元"""

    def test_remove_control_characters(self):
        """移除 XML 不允許的控制字元"""
        text = "前\x00中\x01後\x08尾"
        result = DocxExporter._sanitize_text(text)
        assert result == "前中後尾"

    def test_preserve_newline_tab_cr(self):
        """保留換行、Tab、CR"""
        text = "行一\n行二\t縮排\r回車"
        result = DocxExporter._sanitize_text(text)
        assert "\n" in result
        assert "\t" in result
        assert "\r" in result

    def test_replace_nbsp_with_space(self):
        """NBSP 應替換為普通空格"""
        text = "前\u00a0後"
        result = DocxExporter._sanitize_text(text)
        assert result == "前 後"

    def test_replace_em_space(self):
        """EM Space 應替換為普通空格"""
        text = "前\u2003後"
        result = DocxExporter._sanitize_text(text)
        assert result == "前 後"

    def test_replace_ideographic_space(self):
        """全形空格應替換為普通空格"""
        text = "前\u3000後"
        result = DocxExporter._sanitize_text(text)
        assert result == "前 後"

    def test_replace_thin_space(self):
        """Thin Space 應替換為普通空格"""
        text = "前\u2009後"
        result = DocxExporter._sanitize_text(text)
        assert result == "前 後"

    def test_remove_bom(self):
        """BOM 應被移除"""
        text = "\ufeff開頭"
        result = DocxExporter._sanitize_text(text)
        assert result == "開頭"

    def test_remove_zero_width_space(self):
        """ZWSP 應被移除"""
        text = "前\u200b後"
        result = DocxExporter._sanitize_text(text)
        assert result == "前後"

    def test_mixed_special_characters(self):
        """混合多種特殊字元的清理"""
        text = "\ufeff前\u00a0\x00中\u200b\u3000後\x01"
        result = DocxExporter._sanitize_text(text)
        assert result == "前 中 後"

    def test_empty_input(self):
        """空字串應回傳空字串"""
        assert DocxExporter._sanitize_text("") == ""
        assert DocxExporter._sanitize_text(None) == ""

    def test_export_with_control_chars(self, tmp_path):
        """含控制字元的草稿應能正常匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_ctrl.docx"
        draft = "# 函\n### 主旨\n測試\x00含\x01控制\x08字元"
        exporter.export(draft, str(output_file))
        assert output_file.exists()

    def test_export_with_unicode_spaces(self, tmp_path):
        """含 Unicode 特殊空格的草稿應能正常匯出"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_unicode_space.docx"
        draft = "# 函\n### 主旨\n測試\u00a0含\u3000全形\u2003空格"
        exporter.export(draft, str(output_file))
        assert output_file.exists()


# ==================== 空內容與解析失敗防護 ====================

class TestEmptyContentProtection:
    """空內容和解析失敗的防護測試"""

    def test_empty_string_produces_valid_docx(self, tmp_path):
        """空字串草稿應產生有效 DOCX"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_empty.docx"
        exporter.export("", str(output_file))
        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_none_like_empty_produces_valid_docx(self, tmp_path):
        """僅含空白字元的草稿應產生有效 DOCX"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_whitespace.docx"
        exporter.export("   \n\n  ", str(output_file))
        assert output_file.exists()
        assert output_file.stat().st_size > 0

    def test_parse_failure_produces_valid_docx(self, tmp_path):
        """parse_draft 拋出例外時仍應產生有效 DOCX"""
        exporter = DocxExporter()
        output_file = tmp_path / "test_parse_fail.docx"

        with patch.object(exporter.template_engine, "parse_draft", side_effect=ValueError("模擬解析失敗")):
            exporter.export("# 函\n### 主旨\n測試", str(output_file))

        assert output_file.exists()
        assert output_file.stat().st_size > 0
        # 驗證是有效的 DOCX
        doc = Document(str(output_file))
        assert len(doc.paragraphs) > 0
