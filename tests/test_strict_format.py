"""
嚴格格式模式測試 — 驗證《文書處理手冊》合規性

使用 python-docx 讀取輸出 DOCX，驗證：
1. 字型大小：16/14/12/10 pt
2. 行距：1.5 倍
3. 段前 0.5 行、段後 0
4. 邊距：A4 標準 2.54cm 四邊
5. 自動編號：「一、二、三」及多層級「一、(一)、1.」
"""
import pytest
from docx import Document
from docx.shared import Pt

from src.document.exporter import DocxExporter
from src.core.constants import (
    STRICT_FONT_SIZE_TITLE,
    STRICT_FONT_SIZE_SECTION_LABEL,
    STRICT_FONT_SIZE_BODY,
    STRICT_FONT_SIZE_META,
    STRICT_PAGE_MARGIN,
    FONT_SIZE_TITLE,
)


# ── 輔助函式 ──────────────────────────────────────

def _export_and_read(tmp_path, draft: str, strict: bool = True) -> Document:
    """匯出草稿並讀取 DOCX 回傳 Document 物件。"""
    exporter = DocxExporter(strict_format=strict)
    out = str(tmp_path / "test.docx")
    exporter.export(draft, out)
    return Document(out)


def _find_paragraph_by_text(doc: Document, text: str):
    """找到包含指定文字的段落。"""
    for p in doc.paragraphs:
        if text in p.text:
            return p
    return None


def _get_run_font_size_pt(run) -> float | None:
    """取得 run 的字型大小（pt）。"""
    if run.font.size is not None:
        return run.font.size.pt
    return None


# ── 字型大小驗證 ──────────────────────────────────

class TestStrictFontSizes:
    """嚴格模式字型大小 = 16/14/12/10 pt"""

    DRAFT = (
        "# 函\n\n"
        "**機關**：臺北市政府\n"
        "**受文者**：各級學校\n\n"
        "---\n\n"
        "### 主旨\n測試主旨內容\n\n"
        "### 說明\n測試說明內容\n\n"
        "### 辦法\n請照辦\n"
    )

    def test_title_font_size_16pt(self, tmp_path):
        """公文類型標題應為 16pt"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        p_title = _find_paragraph_by_text(doc, "函")
        assert p_title is not None
        for run in p_title.runs:
            if run.text.strip() == "函":
                assert _get_run_font_size_pt(run) == STRICT_FONT_SIZE_TITLE

    def test_section_label_font_size_14pt(self, tmp_path):
        """段落標籤應為 14pt"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        p_label = _find_paragraph_by_text(doc, "主旨：")
        assert p_label is not None
        for run in p_label.runs:
            if "主旨" in run.text:
                assert _get_run_font_size_pt(run) == STRICT_FONT_SIZE_SECTION_LABEL

    def test_body_font_size_12pt(self, tmp_path):
        """本文內容應為 12pt"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        p_body = _find_paragraph_by_text(doc, "測試主旨內容")
        assert p_body is not None
        for run in p_body.runs:
            if "測試主旨" in run.text:
                assert _get_run_font_size_pt(run) == STRICT_FONT_SIZE_BODY

    def test_meta_font_size_10pt(self, tmp_path):
        """檔頭資訊應為 10pt"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        p_meta = _find_paragraph_by_text(doc, "臺北市政府")
        assert p_meta is not None
        for run in p_meta.runs:
            if "臺北市政府" in run.text:
                assert _get_run_font_size_pt(run) == STRICT_FONT_SIZE_META

    def test_non_strict_uses_old_sizes(self, tmp_path):
        """非嚴格模式應使用舊的 20/16/14/12 pt"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=False)
        p_title = _find_paragraph_by_text(doc, "函")
        assert p_title is not None
        for run in p_title.runs:
            if run.text.strip() == "函":
                assert _get_run_font_size_pt(run) == FONT_SIZE_TITLE


# ── 行距與段距驗證 ──────────────────────────────────

class TestStrictLineSpacing:
    """嚴格模式行距 = 1.5 倍，段前 0.5 行，段後 0"""

    DRAFT = "# 函\n\n### 主旨\n行距測試\n\n### 說明\n說明內容"

    def test_line_spacing_1_5(self, tmp_path):
        """段落行距應為 1.5 倍"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        p = _find_paragraph_by_text(doc, "行距測試")
        assert p is not None
        assert p.paragraph_format.line_spacing == pytest.approx(1.5)

    def test_space_before(self, tmp_path):
        """段前間距應為 body_size × 0.5 = 6pt"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        p = _find_paragraph_by_text(doc, "行距測試")
        assert p is not None
        expected = Pt(STRICT_FONT_SIZE_BODY * 0.5)
        assert p.paragraph_format.space_before == expected

    def test_space_after_zero(self, tmp_path):
        """段後間距應為 0"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        p = _find_paragraph_by_text(doc, "行距測試")
        assert p is not None
        assert p.paragraph_format.space_after == Pt(0)

    def test_non_strict_no_forced_spacing(self, tmp_path):
        """非嚴格模式不應強制設定行距"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=False)
        p = _find_paragraph_by_text(doc, "行距測試")
        assert p is not None
        # 非嚴格模式行距為 None（使用 Word 預設）
        assert p.paragraph_format.line_spacing is None


# ── 邊距驗證 ──────────────────────────────────────

class TestStrictMargins:
    """嚴格模式邊距 = 2.54cm 四邊"""

    DRAFT = "# 函\n\n### 主旨\n邊距測試"

    def _cm_approx(self, emu_value, expected_cm):
        """比較 EMU 值與預期 cm 值（允許 0.01cm 誤差）"""
        actual_cm = emu_value / 360000
        assert abs(actual_cm - expected_cm) < 0.01, (
            f"expected {expected_cm}cm, got {actual_cm:.4f}cm"
        )

    def test_margins_2_54cm(self, tmp_path):
        """嚴格模式四邊邊距應為 2.54cm"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=True)
        section = doc.sections[0]
        self._cm_approx(section.top_margin, STRICT_PAGE_MARGIN)
        self._cm_approx(section.bottom_margin, STRICT_PAGE_MARGIN)
        self._cm_approx(section.left_margin, STRICT_PAGE_MARGIN)
        self._cm_approx(section.right_margin, STRICT_PAGE_MARGIN)

    def test_non_strict_has_wider_side_margins(self, tmp_path):
        """非嚴格模式左右邊距應為 3.17cm"""
        doc = _export_and_read(tmp_path, self.DRAFT, strict=False)
        section = doc.sections[0]
        self._cm_approx(section.left_margin, 3.17)
        self._cm_approx(section.right_margin, 3.17)


# ── 自動編號驗證 ──────────────────────────────────

class TestAutoNumbering:
    """嚴格模式自動編號"""

    def test_multi_line_gets_numbered(self, tmp_path):
        """多行說明應自動加上「一、二、三」"""
        draft = (
            "# 函\n\n### 主旨\n測試\n\n"
            "### 說明\n第一點\n第二點\n第三點\n"
        )
        doc = _export_and_read(tmp_path, draft, strict=True)
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "一、" in full
        assert "二、" in full
        assert "三、" in full

    def test_single_line_no_number(self, tmp_path):
        """單行說明不應加編號"""
        draft = "# 函\n\n### 主旨\n單行測試\n\n### 說明\n只有一行\n"
        doc = _export_and_read(tmp_path, draft, strict=True)
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "一、" not in full

    def test_existing_numbering_preserved(self, tmp_path):
        """已有中文數字編號不應重複編號"""
        draft = (
            "# 函\n\n### 主旨\n測試\n\n"
            "### 說明\n一、已有的第一點\n二、已有的第二點\n"
        )
        doc = _export_and_read(tmp_path, draft, strict=True)
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "一、已有的第一點" in full
        assert "二、已有的第二點" in full
        # 不應出現 "一、一、" 這種重複
        assert "一、一、" not in full

    def test_sub_level_numbering(self):
        """次層級應使用（一）（二）"""
        exporter = DocxExporter(strict_format=True)
        lines = ["第一大點", "  子項目甲", "  子項目乙", "第二大點"]
        result = exporter._auto_number(lines)
        assert any("（一）" in l for l in result)
        assert any("（二）" in l for l in result)

    def test_third_level_numbering(self):
        """第三層級應使用 1. 2."""
        exporter = DocxExporter(strict_format=True)
        lines = ["大項", "  子項", "        細項甲", "        細項乙"]
        result = exporter._auto_number(lines)
        assert any("1." in l for l in result)
        assert any("2." in l for l in result)

    def test_non_strict_no_auto_number(self, tmp_path):
        """非嚴格模式不應自動編號"""
        draft = (
            "# 函\n\n### 主旨\n測試\n\n"
            "### 說明\n第一點\n第二點\n"
        )
        doc = _export_and_read(tmp_path, draft, strict=False)
        full = "\n".join(p.text for p in doc.paragraphs)
        assert "一、" not in full


# ── 向後相容性 ──────────────────────────────────

class TestBackwardsCompatibility:
    """確認預設為嚴格模式，且舊呼叫方式不會壞掉"""

    def test_default_is_strict(self):
        """DocxExporter() 預設為 strict_format=True"""
        exporter = DocxExporter()
        assert exporter.strict_format is True

    def test_old_call_signature_works(self, tmp_path):
        """舊的無參數呼叫方式仍然可用"""
        exporter = DocxExporter()
        out = str(tmp_path / "compat.docx")
        result = exporter.export("# 函\n### 主旨\n測試", out)
        assert result == out
        doc = Document(out)
        assert len(doc.paragraphs) > 0

    def test_explicit_non_strict(self, tmp_path):
        """明確 strict_format=False 使用舊字型大小"""
        exporter = DocxExporter(strict_format=False)
        assert exporter._size_title == FONT_SIZE_TITLE
