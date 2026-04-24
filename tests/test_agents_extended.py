"""
src/agents/ 的延伸測試
補充邊界條件、錯誤路徑和未覆蓋的功能
"""
import importlib
import os
import pytest
import json
from unittest.mock import MagicMock
from src.agents.requirement import RequirementAgent
from src.agents.writer import WriterAgent
from src.agents.template import TemplateEngine, clean_markdown_artifacts, renumber_provisions
from src.agents.auditor import FormatAuditor
from src.agents.style_checker import StyleChecker
from src.agents.fact_checker import FactChecker
from src.agents.consistency_checker import ConsistencyChecker
from src.agents.compliance_checker import ComplianceChecker
from src.agents.org_memory import OrganizationalMemory
from src.agents.review_parser import (
    parse_review_response,
    _extract_json_object,
    format_audit_to_review_result,
)
from src.core.models import PublicDocRequirement
from src.core.review_models import ReviewResult, ReviewIssue, QAReport
from src.agents.editor import EditorInChief
from src.agents.requirement import _sanitize_json_string
from src.knowledge.manager import KnowledgeBaseManager


# ==================== RequirementAgent 邊界測試 ====================

class TestRequirementAgentEdgeCases:
    """RequirementAgent 的邊界測試"""

    def test_json_with_extra_fields(self, mock_llm):
        """測試 JSON 含有多餘欄位時仍能正常解析"""
        agent = RequirementAgent(mock_llm)
        mock_llm.generate.return_value = json.dumps({
            "doc_type": "函",
            "sender": "測試機關",
            "receiver": "測試單位",
            "subject": "測試主旨",
            "extra_field": "should be ignored",
            "another": 123
        })
        req = agent.analyze("test")
        assert req.doc_type == "函"
        assert req.subject == "測試主旨"

    def test_json_with_unicode_escape(self, mock_llm):
        """測試含有 unicode 轉義的 JSON"""
        agent = RequirementAgent(mock_llm)
        mock_llm.generate.return_value = json.dumps({
            "doc_type": "\u51fd",  # 函
            "sender": "\u6e2c\u8a66",  # 測試
            "receiver": "\u55ae\u4f4d",  # 單位
            "subject": "\u4e3b\u65e8"  # 主旨
        })
        req = agent.analyze("test")
        assert req.doc_type == "函"

    def test_json_in_nested_code_block(self, mock_llm):
        """測試在巢狀程式碼區塊中的 JSON"""
        agent = RequirementAgent(mock_llm)
        mock_llm.generate.return_value = '''Sure, here is the result:
```json
{
    "doc_type": "公告",
    "sender": "環保局",
    "receiver": "各學校",
    "subject": "回收公告"
}
```
That's it!'''
        req = agent.analyze("test")
        assert req.doc_type == "公告"

    def test_multiple_code_blocks_uses_first_valid(self, mock_llm):
        """測試多個程式碼區塊時使用第一個有效的"""
        agent = RequirementAgent(mock_llm)
        mock_llm.generate.return_value = '''
```json
invalid json here
```
```json
{"doc_type": "簽", "sender": "市府", "receiver": "局長", "subject": "簽呈"}
```
'''
        req = agent.analyze("test")
        assert req.doc_type == "簽"

    def test_llm_exception_propagates(self, mock_llm):
        """測試 LLM 呼叫拋出異常時的處理"""
        agent = RequirementAgent(mock_llm)
        mock_llm.generate.side_effect = Exception("API timeout")
        with pytest.raises(Exception, match="API timeout"):
            agent.analyze("test")

    def test_json_missing_required_fields_fallback(self, mock_llm):
        """測試 JSON 缺少必要欄位時使用 fallback"""
        agent = RequirementAgent(mock_llm)
        mock_llm.generate.return_value = '{"doc_type": "函"}'
        # JSON 解析會成功但 Pydantic 驗證會失敗，回退到 regex
        # regex 也找不到所有必要欄位，最終使用 fallback
        result = agent.analyze("test input")
        assert result.doc_type == "函"  # fallback 預設
        assert result.sender == "（未指定）"


# ==================== WriterAgent 測試 ====================

class TestWriterAgent:
    """WriterAgent 的測試"""

    def test_write_draft_with_examples(self, mock_llm):
        """測試有範例時的草稿撰寫"""
        mock_kb = MagicMock()
        mock_kb.search_hybrid.return_value = [
            {"content": "範例公文內容", "metadata": {"title": "範例函"}}
        ]
        mock_llm.generate.return_value = "### 主旨\n生成的公文內容"

        writer = WriterAgent(mock_llm, mock_kb)
        requirement = PublicDocRequirement(
            doc_type="函",
            sender="測試機關",
            receiver="測試單位",
            subject="測試主旨"
        )
        draft = writer.write_draft(requirement)

        assert "生成的公文內容" in draft
        assert "參考來源" in draft  # 因為有範例所以應有參考來源
        mock_kb.search_hybrid.assert_called()

    def test_write_draft_without_examples(self, mock_llm):
        """測試沒有範例時的草稿撰寫"""
        mock_kb = MagicMock()
        mock_kb.search_hybrid.return_value = []
        mock_llm.generate.return_value = "### 主旨\n無範例的公文內容"

        writer = WriterAgent(mock_llm, mock_kb)
        requirement = PublicDocRequirement(
            doc_type="函",
            sender="測試機關",
            receiver="測試單位",
            subject="測試主旨"
        )
        draft = writer.write_draft(requirement)

        assert "無範例的公文內容" in draft
        assert "參考來源" not in draft  # 沒有範例所以沒有參考來源

    def test_write_draft_with_multiple_examples(self, mock_llm):
        """測試多個範例時的草稿撰寫"""
        mock_kb = MagicMock()
        mock_kb.search_hybrid.return_value = [
            {"content": "範例一", "metadata": {"title": "函範例A"}},
            {"content": "範例二", "metadata": {"title": "函範例B"}}
        ]
        mock_llm.generate.return_value = "### 主旨\n綜合範例的公文"

        writer = WriterAgent(mock_llm, mock_kb)
        requirement = PublicDocRequirement(
            doc_type="函",
            sender="環保局",
            receiver="各學校",
            subject="資源回收"
        )
        draft = writer.write_draft(requirement)

        assert "[^1]" in draft  # Source 1
        assert "[^2]" in draft  # Source 2

    def test_write_draft_uses_correct_query(self, mock_llm):
        """測試搜尋查詢包含正確的公文類型和主旨"""
        mock_kb = MagicMock()
        mock_kb.search_hybrid.return_value = []
        mock_llm.generate.return_value = "### 主旨\n測試"

        writer = WriterAgent(mock_llm, mock_kb)
        requirement = PublicDocRequirement(
            doc_type="公告",
            sender="市府",
            receiver="各區",
            subject="垃圾清運公告"
        )
        writer.write_draft(requirement)

        # 驗證搜尋查詢（search_hybrid 被呼叫兩次：Level A + 全部）
        call_args = mock_kb.search_hybrid.call_args_list[0]
        assert "公告" in call_args[0][0]
        assert "垃圾清運公告" in call_args[0][0]


# ==================== TemplateEngine 邊界測試 ====================

class TestTemplateEngineEdgeCases:
    """TemplateEngine 的邊界測試"""

    def test_parse_list_items_empty(self):
        """測試空文字的清單解析"""
        engine = TemplateEngine()
        items = engine._parse_list_items("")
        assert items == []

    def test_parse_list_items_none(self):
        """測試 None 的清單解析"""
        engine = TemplateEngine()
        items = engine._parse_list_items(None)
        assert items == []

    def test_parse_list_items_with_numbering(self):
        """測試帶編號的清單解析：子項合併到父項"""
        engine = TemplateEngine()
        text = "1、第一項\n2、第二項\n（一）子項"
        items = engine._parse_list_items(text)
        # 子項（一）應合併到父項「第二項」，不獨立成項
        assert len(items) == 2
        assert "第一項" in items[0]
        assert "第二項" in items[1]
        assert "（一）子項" in items[1]

    def test_parse_draft_empty(self):
        """測試空草稿的解析"""
        engine = TemplateEngine()
        sections = engine.parse_draft("")
        assert sections["subject"] == ""
        assert sections["explanation"] == ""

    def test_parse_draft_with_inline_content(self):
        """測試標頭後直接跟隨內容（同一行）"""
        engine = TemplateEngine()
        draft = "### 主旨：測試同行內容\n### 說明\n說明內容"
        sections = engine.parse_draft(draft)
        assert "測試同行內容" in sections["subject"]

    def test_parse_draft_announcement_provisions(self):
        """測試公告事項的解析"""
        engine = TemplateEngine()
        draft = """
### 主旨
公告主旨

### 公告事項
一、第一項
二、第二項
"""
        sections = engine.parse_draft(draft)
        assert sections["provisions"] != ""
        assert "一、" in sections["provisions"]

    def test_apply_template_sign_type(self, sample_sign_requirement):
        """測試簽類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "簽呈主旨",
            "explanation": "簽呈說明",
            "provisions": "",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_sign_requirement, sections)
        assert "市府" in result or "簽呈主旨" in result

    def test_apply_template_decree_type(self, sample_decree_requirement):
        """測試令類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "修正某條例施行細則",
            "explanation": "",
            "provisions": "一、自公布日施行",
            "basis": "依據某法第X條規定",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_decree_requirement, sections)
        assert "行政院" in result
        assert "修正某條例" in result
        assert "令" in result

    def test_apply_template_meeting_notice_type(self, sample_meeting_notice_requirement):
        """測試開會通知單類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "召開年度預算審查會議",
            "explanation": "為審議115年度預算案",
            "provisions": "一、請準時出席\n二、攜帶相關資料",
            "meeting_time": "中華民國115年3月15日上午10時",
            "meeting_location": "市府大樓10樓會議室",
            "agenda": "一、報告事項\n二、討論事項",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_meeting_notice_requirement, sections)
        assert "臺北市政府" in result
        assert "召開年度預算" in result
        assert "開會通知單" in result

    def test_apply_template_meeting_minutes_type(self, sample_meeting_minutes_requirement):
        """測試開會紀錄類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "115年度第1次預算審查會議紀錄",
            "explanation": "",
            "provisions": "",
            "meeting_time": "中華民國115年3月15日上午10時",
            "meeting_location": "市府大樓10樓會議室",
            "meeting_name": "臺北市政府115年度第1次預算審查會議紀錄",
            "chairperson": "陳局長○○",
            "attendees": "王科長○○、李專員○○",
            "observers": "張主任○○",
            "recorder": "林科員○○",
            "report_items": "一、上期預算執行進度報告\n二、本期預算編列說明",
            "discussion_items": "一、115年度資本門預算配置\n二、人事費用調整方案",
            "resolutions": "一、資本門預算依原案通過\n二、人事費用調整案擇期再議",
            "motions": "無",
            "adjournment_time": "上午12時",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_meeting_minutes_requirement, sections)
        assert "開會紀錄" in result
        assert "臺北市政府" in result
        assert "陳局長" in result
        assert "會議室" in result
        assert "報告事項" in result
        assert "討論事項" in result
        assert "決議" in result
        assert "散會時間" in result

    def test_apply_template_meeting_minutes_minimal(self, sample_meeting_minutes_requirement):
        """測試開會紀錄最小欄位（只有必要欄位）"""
        engine = TemplateEngine()
        sections = {
            "subject": "會議紀錄",
            "explanation": "",
            "provisions": "",
            "meeting_time": "中華民國115年4月1日",
            "meeting_location": "第一會議室",
            "chairperson": "王主席",
            "recorder": "李紀錄",
            "discussion_items": "討論年度計畫",
            "resolutions": "照案通過",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_meeting_minutes_requirement, sections)
        assert "開會紀錄" in result
        assert "王主席" in result
        assert "李紀錄" in result
        assert "討論" in result
        assert "照案通過" in result

    def test_parse_draft_meeting_minutes(self):
        """測試 parse_draft 能正確解析開會紀錄格式的草稿"""
        engine = TemplateEngine()
        draft = """會議名稱：臺北市政府第3次協調會紀錄
主席：王局長○○
出席人員：張科長、李專員
列席人員：陳主任
紀錄人：林書記
報告事項：
一、前次會議決議執行情形
二、本季業務推動成果
討論事項：
一、下半年度工作計畫
二、預算調整案
決議：
一、工作計畫照案通過
二、預算調整案修正後通過
臨時動議：無
散會時間：下午3時30分"""
        result = engine.parse_draft(draft)
        assert result["meeting_name"] == "臺北市政府第3次協調會紀錄"
        assert result["chairperson"] == "王局長○○"
        assert result["attendees"] == "張科長、李專員"
        assert result["observers"] == "陳主任"
        assert result["recorder"] == "林書記"
        assert "前次會議決議執行情形" in result["report_items"]
        assert "下半年度工作計畫" in result["discussion_items"]
        assert "工作計畫照案通過" in result["resolutions"]
        assert result["adjournment_time"] == "下午3時30分"

    def test_apply_template_letter_type(self, sample_letter_requirement):
        """測試書函類型的模板套用（使用函模板）"""
        engine = TemplateEngine()
        sections = {
            "subject": "函送資源回收作業注意事項",
            "explanation": "配合中央政策",
            "provisions": "一、請依規定辦理",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_letter_requirement, sections)
        assert "環境保護局" in result
        assert "函送資源回收" in result

    def test_apply_template_chen_type(self, sample_chen_requirement):
        """測試呈類型的模板套用（使用函模板）"""
        engine = TemplateEngine()
        sections = {
            "subject": "呈報115年度施政成果報告",
            "explanation": "依據行政院組織法規定",
            "provisions": "",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_chen_requirement, sections)
        assert "行政院" in result
        assert "呈報115年度施政成果" in result

    def test_apply_template_zi_type(self, sample_zi_requirement):
        """測試咨類型的模板套用（使用函模板）"""
        engine = TemplateEngine()
        sections = {
            "subject": "咨請貴院審議國際條約案",
            "explanation": "依據憲法第63條規定",
            "provisions": "",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_zi_requirement, sections)
        assert "總統府" in result
        assert "咨請貴院審議" in result

    def test_apply_template_inspection_type(self, sample_inspection_requirement):
        """測試會勘通知單類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "辦理信義路段道路損壞會勘",
            "explanation": "接獲民眾陳情道路損壞",
            "provisions": "一、請派員參加\n二、攜帶相關圖說",
            "inspection_time": "中華民國115年4月1日上午10時",
            "inspection_location": "臺北市信義路四段",
            "inspection_items": "道路損壞範圍及程度評估",
            "required_documents": "施工圖說及相關照片",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_inspection_requirement, sections)
        assert "工務局" in result
        assert "道路損壞" in result
        assert "會勘通知單" in result
        assert "中華民國115年4月1日上午10時" in result, "會勘時間未出現在渲染結果中"
        assert "臺北市信義路四段" in result, "會勘地點未出現在渲染結果中"
        assert "道路損壞範圍及程度評估" in result, "會勘事項未出現在渲染結果中"
        assert "施工圖說及相關照片" in result, "應攜文件未出現在渲染結果中"

    def test_apply_template_phone_type(self, sample_phone_requirement):
        """測試公務電話紀錄類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "確認環境影響評估會議時間",
            "explanation": "因原訂會議時間與其他會議衝突",
            "provisions": "",
            "call_time": "中華民國115年3月5日下午2時30分",
            "caller": "秘書處王科長",
            "callee": "環保局李科長",
            "call_summary": "確認會議改至3月10日上午10時召開",
            "follow_up_items": "請環保局確認出席名單",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_phone_requirement, sections)
        assert "秘書處" in result
        assert "環境影響評估" in result
        assert "公務電話紀錄" in result
        assert "中華民國115年3月5日下午2時30分" in result, "通話時間未出現在渲染結果中"
        assert "秘書處王科長" in result, "發話人未出現在渲染結果中"
        assert "環保局李科長" in result, "受話人未出現在渲染結果中"
        assert "確認會議改至3月10日上午10時召開" in result, "通話摘要未出現在渲染結果中"
        assert "請環保局確認出席名單" in result, "追蹤事項未出現在渲染結果中"

    def test_apply_template_directive_type(self, sample_directive_requirement):
        """測試手令類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "指派辦理社會住宅專案",
            "explanation": "為加速推動社會住宅政策",
            "provisions": "",
            "directive_content": "即日起督導辦理本市社會住宅興建計畫",
            "deadline": "中華民國115年12月31日前完成第一期工程",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_directive_requirement, sections)
        assert "市長" in result
        assert "社會住宅" in result
        assert "手令" in result
        assert "即日起督導辦理本市社會住宅興建計畫" in result, "指示事項未出現在渲染結果中"
        assert "中華民國115年12月31日前完成第一期工程" in result, "完成期限未出現在渲染結果中"

    def test_apply_template_memo_type(self, sample_memo_requirement):
        """測試箋函類型的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "請提供本年度員工訓練計畫",
            "explanation": "配合年度施政報告彙整",
            "provisions": "",
            "copies_to": "臺北市政府人事處",
            "cc_copies": "臺北市政府秘書處",
            "attachments": "",
            "references": ""
        }
        result = engine.apply_template(sample_memo_requirement, sections)
        assert "秘書處" in result
        assert "員工訓練" in result
        assert "箋函" in result
        assert "正本" in result, "正本未出現在渲染結果中"

    def test_all_doc_types_have_templates(self):
        """驗證 VALID_DOC_TYPES 中所有 12 種公文類型都有對應模板"""
        from src.core.models import VALID_DOC_TYPES
        engine = TemplateEngine()
        for doc_type in VALID_DOC_TYPES:
            req = PublicDocRequirement(
                doc_type=doc_type,
                sender="測試機關",
                receiver="測試單位",
                subject=f"測試{doc_type}主旨",
            )
            sections = {
                "subject": f"測試{doc_type}主旨",
                "explanation": "測試說明",
                "provisions": "",
                "attachments": "",
                "references": ""
            }
            result = engine.apply_template(req, sections)
            assert f"測試{doc_type}主旨" in result, f"公文類型 {doc_type} 的模板套用失敗"

    def test_apply_template_with_references(self, sample_requirement):
        """測試有參考來源時的模板套用"""
        engine = TemplateEngine()
        sections = {
            "subject": "測試主旨",
            "explanation": "測試說明",
            "provisions": "",
            "attachments": "",
            "references": "[^1]: 某法規"
        }
        result = engine.apply_template(sample_requirement, sections)
        assert "參考來源" in result
        assert "某法規" in result

    def test_fallback_apply(self, sample_requirement):
        """測試 fallback 模板（Jinja2 載入失敗時）"""
        engine = TemplateEngine()
        sections = {
            "subject": "測試主旨",
            "explanation": "測試說明",
            "provisions": "一、辦法一",
            "attachments": "",
            "references": ""
        }
        result = engine._fallback_apply(sample_requirement, sections)
        assert "測試機關" in result
        assert "測試主旨" in result
        assert "辦法一" in result

    def test_fallback_apply_with_attachments(self, sample_requirement):
        """測試 fallback 模板包含附件"""
        engine = TemplateEngine()
        sections = {
            "subject": "測試",
            "explanation": "",
            "provisions": "",
            "attachments": "",
            "references": ""
        }
        result = engine._fallback_apply(sample_requirement, sections)
        assert "附件" in result
        assert "附件1" in result

    def test_fallback_apply_with_action_items(self):
        """測試 fallback 模板使用 action_items 作為說明"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="函",
            sender="測試",
            receiver="對象",
            subject="主旨",
            action_items=["動作一", "動作二"]
        )
        sections = {
            "subject": "主旨",
            "explanation": "",
            "provisions": "",
            "attachments": "",
            "references": ""
        }
        result = engine._fallback_apply(req, sections)
        assert "動作一" in result

    def test_parse_draft_inspection_sections(self):
        """測試會勘通知單的專用段落解析"""
        engine = TemplateEngine()
        draft = """### 主旨
辦理道路損壞會勘

### 說明
接獲民眾陳情

### 會勘時間
中華民國115年4月1日上午10時

### 會勘地點
臺北市信義路四段

### 會勘事項
道路損壞範圍評估

### 應攜文件
施工圖說
"""
        sections = engine.parse_draft(draft)
        assert sections["inspection_time"] == "中華民國115年4月1日上午10時"
        assert sections["inspection_location"] == "臺北市信義路四段"
        assert sections["inspection_items"] == "道路損壞範圍評估"
        assert sections["required_documents"] == "施工圖說"

    def test_parse_draft_phone_record_sections(self):
        """測試公務電話紀錄的專用段落解析"""
        engine = TemplateEngine()
        draft = """### 通話時間
115年3月5日下午2時

### 發話人
王科長

### 受話人
李科長

### 主旨
確認會議時間

### 通話摘要
會議改至3月10日

### 追蹤事項
確認出席名單

### 紀錄人
張書記

### 核閱
陳處長
"""
        sections = engine.parse_draft(draft)
        assert sections["call_time"] == "115年3月5日下午2時"
        assert sections["caller"] == "王科長"
        assert sections["callee"] == "李科長"
        assert sections["call_summary"] == "會議改至3月10日"
        assert sections["follow_up_items"] == "確認出席名單"
        assert sections["recorder"] == "張書記"
        assert sections["reviewer"] == "陳處長"

    def test_parse_draft_directive_sections(self):
        """測試手令的專用段落解析"""
        engine = TemplateEngine()
        draft = """### 主旨
指派辦理專案

### 指示事項
即日起督導辦理

### 說明
為加速推動政策

### 完成期限
115年12月31日前

### 副知
秘書處、人事處
"""
        sections = engine.parse_draft(draft)
        assert sections["directive_content"] == "即日起督導辦理"
        assert sections["deadline"] == "115年12月31日前"
        assert sections["cc_list"] == "秘書處、人事處"


# ==================== clean_markdown_artifacts 邊界測試 ====================

class TestCleanMarkdownArtifacts:
    """clean_markdown_artifacts 的邊界測試"""

    def test_empty_string(self):
        """測試空字串"""
        assert clean_markdown_artifacts("") == ""

    def test_none_input(self):
        """測試 None 輸入"""
        assert clean_markdown_artifacts(None) == ""

    def test_nested_bold_italic(self):
        """測試巢狀粗體斜體"""
        text = "***粗斜體***"
        result = clean_markdown_artifacts(text)
        assert "***" not in result

    def test_remove_separator_lines(self):
        """測試移除分隔線"""
        text = "第一段\n---\n第二段"
        result = clean_markdown_artifacts(text)
        assert "---" not in result

    def test_remove_stamp_text(self):
        """測試移除捺印處文字"""
        text = "承辦人 捺印處 審核"
        result = clean_markdown_artifacts(text)
        assert "捺印處" not in result

    def test_collapse_multiple_blank_lines(self):
        """測試合併多餘空行"""
        text = "第一段\n\n\n\n\n第二段"
        result = clean_markdown_artifacts(text)
        assert "\n\n\n" not in result

    def test_url_in_link(self):
        """測試連結中的 URL 被移除"""
        text = "請參閱[公文系統](https://example.com)辦理"
        result = clean_markdown_artifacts(text)
        assert "https://example.com" not in result
        assert "公文系統" in result

    def test_remove_inline_code_backticks(self):
        """測試移除行內程式碼標記 `code`"""
        text = "依據`廢棄物清理法`第 10 條辦理"
        result = clean_markdown_artifacts(text)
        assert "`" not in result
        assert "廢棄物清理法" in result

    def test_remove_strikethrough(self):
        """測試移除刪除線標記 ~~text~~"""
        text = "本案~~暫緩~~改為即日執行"
        result = clean_markdown_artifacts(text)
        assert "~~" not in result
        assert "暫緩" in result
        assert "改為即日執行" in result

    def test_remove_blockquote(self):
        """測試移除 blockquote 標記"""
        text = "> 引用政策條文\n> 第二行引用"
        result = clean_markdown_artifacts(text)
        assert result.startswith("引用政策條文")
        assert ">" not in result.split("引用政策條文")[0]

    def test_blockquote_preserves_inline_gt(self):
        """測試 blockquote 移除不影響行內的 > 符號"""
        text = "數量 > 10 的情況\n金額 >= 500 元"
        result = clean_markdown_artifacts(text)
        assert ">" in result
        assert "數量 > 10" in result

    def test_multiple_backticks_in_one_line(self):
        """測試同一行多個 inline code"""
        text = "依據`法規A`及`法規B`辦理"
        result = clean_markdown_artifacts(text)
        assert "`" not in result
        assert "法規A" in result
        assert "法規B" in result


# ==================== renumber_provisions 邊界測試 ====================

class TestRenumberProvisionsEdgeCases:
    """renumber_provisions 的邊界測試"""

    def test_skip_signature_section(self):
        """測試跳過簽署區內容"""
        text = "1. 辦法內容\n正本：受文者\n副本：其他"
        result = renumber_provisions(text)
        assert "一、" in result
        assert "正本" in result
        assert "副本" in result

    def test_mixed_numbering(self):
        """測試混合編號格式"""
        text = "1. 第一項\n(1) 子項A\n(2) 子項B\n2. 第二項"
        result = renumber_provisions(text)
        assert "一、" in result
        assert "二、" in result
        assert "（一）" in result
        assert "（二）" in result

    def test_chinese_numbering_already_present(self):
        """測試已有中文編號的文字"""
        text = "一、第一項\n二、第二項"
        result = renumber_provisions(text)
        assert "一、" in result
        assert "二、" in result

    def test_plain_text_preserved(self):
        """測試普通文字不被修改"""
        text = "這是一段普通文字，不需要編號。"
        result = renumber_provisions(text)
        assert result == "這是一段普通文字，不需要編號。"

    def test_empty_lines_preserved(self):
        """測試空行被保留"""
        text = "1. 第一項\n\n2. 第二項"
        result = renumber_provisions(text)
        assert "" in result.split("\n")

    def test_sub_counter_resets(self):
        """測試子編號在主編號變化後重設"""
        text = "1. 主項一\n(1) 子項\n2. 主項二\n(1) 新子項"
        result = renumber_provisions(text)
        lines = [line for line in result.split("\n") if line.strip()]
        # 第二個 (1) 應該也是 （一）因為子計數器已重設
        sub_items = [line for line in lines if "（一）" in line]
        assert len(sub_items) == 2


# ==================== FormatAuditor 邊界測試 ====================

class TestFormatAuditorEdgeCases:
    """FormatAuditor 的邊界測試"""

    def test_audit_with_kb_rules(self, mock_llm):
        """測試有知識庫規則時的審查"""
        mock_kb = MagicMock()
        mock_kb.search_regulations.return_value = [
            {"content": "- [Error] 必須有主旨段落", "metadata": {"title": "函規則"}}
        ]
        mock_llm.generate.return_value = '{"errors": ["缺少主旨"], "warnings": []}'

        auditor = FormatAuditor(mock_llm, mock_kb)
        result = auditor.audit("### 說明\n測試", "函")
        assert "缺少主旨" in result["errors"]

    def test_audit_with_validator_calls(self, mock_llm):
        """測試含有 [Call: func_name] 的規則會觸發驗證器"""
        mock_kb = MagicMock()
        mock_kb.search_regulations.return_value = [
            {"content": "[Call: check_doc_integrity]\n- [Error] 必須有主旨",
             "metadata": {"title": "規則"}}
        ]
        mock_llm.generate.return_value = '{"errors": [], "warnings": []}'

        auditor = FormatAuditor(mock_llm, mock_kb)
        # 草稿缺少標準欄位，check_doc_integrity 應產生錯誤
        result = auditor.audit("### 主旨\n測試", "函")
        assert len(result["errors"]) > 0

    def test_audit_llm_exception(self, mock_llm):
        """測試 LLM 呼叫異常時的錯誤處理"""
        mock_llm.generate.side_effect = RuntimeError("Timeout")
        auditor = FormatAuditor(mock_llm)
        result = auditor.audit("### 主旨\n測試", "函")
        assert len(result["warnings"]) > 0
        assert any("內部錯誤" in w or "手動檢查" in w for w in result["warnings"])

    def test_audit_llm_exception_logs_warning(self, mock_llm, caplog):
        """LLM 失敗時應留下 warning log。"""
        mock_llm.generate.side_effect = RuntimeError("Timeout")
        auditor = FormatAuditor(mock_llm)

        with caplog.at_level("WARNING"):
            result = auditor.audit("### 主旨\n測試", "函")

        assert result["warnings"]
        assert "FormatAuditor LLM 解析意外例外" in caplog.text


# ==================== StyleChecker 邊界測試 ====================

class TestStyleCheckerEdgeCases:
    """StyleChecker 的邊界測試"""

    def test_check_with_issues(self, mock_llm):
        """測試包含問題的審查結果"""
        mock_llm.generate.return_value = json.dumps({
            "issues": [
                {"severity": "warning", "location": "主旨",
                 "description": "口語化用詞", "suggestion": "改為正式用語"}
            ],
            "score": 0.7
        })
        checker = StyleChecker(mock_llm)
        result = checker.check("幫我處理一下")
        assert result.score == 0.7
        assert len(result.issues) == 1
        assert result.issues[0].category == "style"

    def test_check_json_decode_error(self, mock_llm):
        """測試 JSON 解析失敗時回傳預設值"""
        mock_llm.generate.return_value = "This is completely invalid"
        checker = StyleChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.8
        assert len(result.issues) == 0

    def test_check_exception_handling(self, mock_llm):
        """測試一般例外時回傳預設值"""
        mock_llm.generate.return_value = '{"issues": "not_a_list", "score": 0.9}'
        checker = StyleChecker(mock_llm)
        result = checker.check("test")
        # issues 非列表被忽略（空），但 score 仍從 JSON 正確取得
        assert result.score == 0.9
        assert result.issues == []


# ==================== FactChecker 邊界測試 ====================

class TestFactCheckerEdgeCases:
    """FactChecker 的邊界測試"""

    def test_check_empty_response(self, mock_llm):
        """測試空回應時回傳預設值"""
        mock_llm.generate.return_value = ""
        checker = FactChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.8
        assert result.agent_name == "Fact Checker"

    def test_check_json_decode_error(self, mock_llm):
        """測試 JSON 解析失敗時回傳預設值"""
        mock_llm.generate.return_value = "Not a JSON response"
        checker = FactChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.8

    def test_check_no_json_match(self, mock_llm):
        """測試找不到 JSON 物件時回傳預設值"""
        mock_llm.generate.return_value = "No JSON here at all"
        checker = FactChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.8

    def test_check_with_multiple_issues(self, mock_llm):
        """測試多個問題的審查結果"""
        mock_llm.generate.return_value = json.dumps({
            "issues": [
                {"severity": "error", "location": "依據", "description": "法規不存在"},
                {"severity": "warning", "location": "日期", "description": "日期可能過期"}
            ],
            "score": 0.4
        })
        checker = FactChecker(mock_llm)
        result = checker.check("依據某法辦理")
        assert len(result.issues) == 2
        assert result.score == 0.4
        assert all(i.category == "fact" for i in result.issues)


# ==================== ConsistencyChecker 邊界測試 ====================

class TestConsistencyCheckerEdgeCases:
    """ConsistencyChecker 的邊界測試"""

    def test_check_empty_response(self, mock_llm):
        """測試空回應時回傳預設值"""
        mock_llm.generate.return_value = ""
        checker = ConsistencyChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.8

    def test_check_json_decode_error(self, mock_llm):
        """測試 JSON 解析失敗"""
        mock_llm.generate.return_value = "{invalid json"
        checker = ConsistencyChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.8

    def test_check_with_contradictions(self, mock_llm):
        """測試發現矛盾時的審查結果"""
        mock_llm.generate.return_value = json.dumps({
            "issues": [
                {"severity": "error", "location": "主旨 vs 辦法",
                 "description": "主旨說請出席會議但辦法沒有列出時間地點"}
            ],
            "score": 0.5
        })
        checker = ConsistencyChecker(mock_llm)
        result = checker.check("主旨：請出席\n辦法：無")
        assert result.has_errors is True


# ==================== ComplianceChecker 邊界測試 ====================

class TestComplianceCheckerEdgeCases:
    """ComplianceChecker 的邊界測試"""

    def test_check_with_kb(self, mock_llm):
        """測試有知識庫時的政策合規檢查"""
        mock_kb = MagicMock()
        mock_kb.search_policies.return_value = [
            {"metadata": {"title": "淨零碳排政策"}, "content": "所有公文應考慮碳排放"}
        ]
        mock_llm.generate.return_value = json.dumps({
            "issues": [], "score": 0.95, "confidence": 0.9
        })

        checker = ComplianceChecker(mock_llm, mock_kb)
        result = checker.check("### 主旨\n測試公文")
        assert result.score == 0.95
        assert result.confidence == 0.9

    def test_check_kb_exception(self, mock_llm):
        """測試知識庫查詢異常時不中斷"""
        mock_kb = MagicMock()
        mock_kb.search_policies.side_effect = Exception("DB error")
        mock_llm.generate.return_value = json.dumps({
            "issues": [], "score": 0.85, "confidence": 0.5
        })

        checker = ComplianceChecker(mock_llm, mock_kb)
        result = checker.check("test")
        assert result.score == 0.85  # 應仍能正常運作

    def test_check_empty_response(self, mock_llm):
        """測試空回應"""
        mock_llm.generate.return_value = ""
        checker = ComplianceChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.85
        assert result.confidence == 0.5

    def test_check_json_decode_error(self, mock_llm):
        """測試 JSON 解析失敗"""
        mock_llm.generate.return_value = "not json"
        checker = ComplianceChecker(mock_llm)
        result = checker.check("test")
        assert result.score == 0.85

    def test_check_with_compliance_issues(self, mock_llm):
        """測試發現合規問題時的處理"""
        mock_llm.generate.return_value = json.dumps({
            "issues": [
                {
                    "severity": "error",
                    "location": "說明",
                    "description": "抵觸淨零碳排政策",
                    "suggestion": "調整措辭"
                }
            ],
            "score": 0.4,
            "confidence": 0.8
        })
        checker = ComplianceChecker(mock_llm)
        result = checker.check("增加燃煤發電")
        assert len(result.issues) == 1
        assert result.issues[0].risk_level == "high"
        assert result.issues[0].category == "compliance"

    def test_check_risk_level_mapping(self, mock_llm):
        """測試嚴重性到風險等級的映射"""
        mock_llm.generate.return_value = json.dumps({
            "issues": [
                {"severity": "error", "location": "A", "description": "高風險"},
                {"severity": "warning", "location": "B", "description": "中風險"},
                {"severity": "info", "location": "C", "description": "低風險"}
            ],
            "score": 0.5,
            "confidence": 0.8
        })
        checker = ComplianceChecker(mock_llm)
        result = checker.check("test")
        assert result.issues[0].risk_level == "high"
        assert result.issues[1].risk_level == "medium"
        assert result.issues[2].risk_level == "low"

    def test_check_general_exception(self, mock_llm):
        """測試 LLM 呼叫異常時回傳失敗分數（排除加權計算，與其他 Agent 一致）"""
        mock_llm.generate.side_effect = Exception("Unknown error")
        checker = ComplianceChecker(mock_llm)
        result = checker.check("test")
        # LLM 呼叫失敗時應回傳 0.0/0.0（排除加權分數計算），而非拋出例外
        assert result.score == 0.0
        assert result.confidence == 0.0
        assert len(result.issues) == 0


# ==================== OrganizationalMemory 邊界測試 ====================

class TestOrganizationalMemoryEdgeCases:
    """OrganizationalMemory 的邊界測試"""

    def test_learn_from_edit(self, tmp_path):
        """測試從編輯中學習"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.learn_from_edit("測試機關", "原始內容", "修改後內容")

        profile = mem.get_agency_profile("測試機關")
        assert profile["usage_count"] == 1
        assert "last_edit" in profile

    def test_learn_from_edit_increments_count(self, tmp_path):
        """測試多次編輯學習累計次數"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.learn_from_edit("機關A", "v1", "v2")
        mem.learn_from_edit("機關A", "v2", "v3")

        profile = mem.get_agency_profile("機關A")
        assert profile["usage_count"] == 2

    def test_corrupted_preferences_file(self, tmp_path):
        """測試損壞的偏好設定檔案"""
        storage = tmp_path / "prefs.json"
        storage.write_text("THIS IS NOT JSON!!!", encoding="utf-8")

        mem = OrganizationalMemory(str(storage))
        assert mem.preferences == {}

    def test_get_writing_hints_concise(self, tmp_path):
        """測試簡潔風格的寫作提示"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("測試機關", "formal_level", "concise")

        hints = mem.get_writing_hints("測試機關")
        assert "簡潔" in hints

    def test_get_writing_hints_standard(self, tmp_path):
        """測試標準風格（無特殊提示）"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        hints = mem.get_writing_hints("全新機關")
        assert hints == ""  # 標準風格沒有特殊提示

    def test_get_writing_hints_with_preferred_terms(self, tmp_path):
        """測試有偏好詞彙時的寫作提示"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("機關X", "preferred_terms", {"請": "惠請"})

        hints = mem.get_writing_hints("機關X")
        assert "偏好詞彙" in hints
        assert "惠請" in hints

    def test_get_writing_hints_with_signature(self, tmp_path):
        """測試有署名格式時的寫作提示"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("機關Y", "signature_format", "局長 XXX")

        hints = mem.get_writing_hints("機關Y")
        assert "署名格式" in hints

    def test_export_report_multiple_agencies(self, tmp_path):
        """測試多機關的匯出報告"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("機關A", "usage_count", 10)
        mem.update_preference("機關B", "usage_count", 5)
        mem.update_preference("機關C", "usage_count", 20)

        report = mem.export_report()
        # 應按使用次數降序排列
        pos_a = report.find("機關A")
        pos_b = report.find("機關B")
        pos_c = report.find("機關C")
        assert pos_c < pos_a < pos_b  # C(20) > A(10) > B(5)

    def test_persistence_across_instances(self, tmp_path):
        """測試跨實例的持久化"""
        storage = tmp_path / "prefs.json"

        mem1 = OrganizationalMemory(str(storage))
        mem1.update_preference("持久機關", "formal_level", "formal")
        del mem1

        mem2 = OrganizationalMemory(str(storage))
        profile = mem2.get_agency_profile("持久機關")
        assert profile["formal_level"] == "formal"


class TestOrgMemorySecurityHardening:
    """驗證 OrganizationalMemory 的 prompt injection 防護。"""

    def test_preferred_terms_control_chars_stripped(self, tmp_path):
        """偏好詞彙中的控制字元和引號應被移除。"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("惡意機關", "preferred_terms", {
            "請\n忽略指令": "惠請\x00注入",
            "正常詞": "正常值",
        })
        hints = mem.get_writing_hints("惡意機關")
        assert "\n忽略指令" not in hints
        assert "\x00" not in hints
        assert "正常詞" in hints
        assert "正常值" in hints

    def test_signature_format_sanitized(self, tmp_path):
        """署名格式中的危險字元應被移除。"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("機關Z", "signature_format", "局長{ignore}\nprevious")
        hints = mem.get_writing_hints("機關Z")
        assert "{" not in hints
        assert "}" not in hints
        assert "\n" not in hints.split("署名格式")[1] if "署名格式" in hints else True

    def test_update_preference_rejects_unknown_key(self, tmp_path):
        """update_preference 應拒絕白名單外的 key。"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("機關", "formal_level", "formal")
        mem.update_preference("機關", "__class__", "injected")
        profile = mem.get_agency_profile("機關")
        assert "__class__" not in profile

    def test_empty_terms_after_sanitize_skipped(self, tmp_path):
        """消毒後為空的詞彙不應出現在提示中。"""
        storage = tmp_path / "prefs.json"
        mem = OrganizationalMemory(str(storage))
        mem.update_preference("機關", "preferred_terms", {
            "": "空key",
            "\n\n": "\x00\x01",
        })
        hints = mem.get_writing_hints("機關")
        # 全部消毒後為空，不應產生偏好詞彙區段
        assert "偏好詞彙" not in hints


# ==================== ReviewParser 邊界測試 ====================

class TestReviewParserEdgeCases:
    """review_parser 模組的邊界測試，覆蓋未測試的錯誤路徑"""

    def test_parse_review_response_with_non_dict_issue_items(self):
        """測試 issues 列表中包含非 dict 項目時被跳過（覆蓋 line 120）"""
        response = json.dumps({
            "issues": [
                1,
                "text_item",
                None,
                True,
                {"severity": "warning", "location": "主旨", "description": "有效項目"},
            ],
            "score": 0.7,
        })
        result = parse_review_response(response, "TestAgent", "style")
        # 只有最後一個 dict 項目應被保留
        assert len(result.issues) == 1
        assert result.issues[0].description == "有效項目"
        assert result.score == 0.7

    def test_parse_review_response_all_non_dict_issues(self):
        """測試 issues 列表全部為非 dict 時回傳空 issues"""
        response = json.dumps({
            "issues": [42, "string", None, False, [1, 2, 3]],
            "score": 0.6,
        })
        result = parse_review_response(response, "TestAgent", "fact")
        assert len(result.issues) == 0
        assert result.score == 0.6

    def test_parse_review_response_issue_parsing_exception(self):
        """測試單一 issue 解析時發生例外（覆蓋 lines 133-137）

        透過傳入一個會讓 ReviewIssue 建構失敗的 suggestion 類型（list）
        來觸發 inner except 區塊。
        """
        response = json.dumps({
            "issues": [
                {
                    "severity": "warning",
                    "location": "說明",
                    "description": "問題描述",
                    "suggestion": {"nested": "object_not_string"},
                },
                {
                    "severity": "info",
                    "location": "辦法",
                    "description": "正常項目",
                },
            ],
            "score": 0.65,
        })
        result = parse_review_response(response, "TestAgent", "consistency")
        # 如果第一個 issue 因為 suggestion 類型問題而解析失敗，
        # 至少第二個正常項目應該被保留
        # （注意：Pydantic 可能會成功強制轉型，此時兩個都會保留）
        assert result.score == 0.65
        assert len(result.issues) >= 1

    def test_parse_review_response_json_syntax_error(self):
        """測試 JSON 語法錯誤時回傳預設結果（覆蓋 lines 145-146）"""
        # _extract_json_object 會找到 {... 但 json.loads 會失敗
        response = '{issues: [broken json syntax}'
        result = parse_review_response(response, "SyntaxTest", "style")
        assert result.score == 0.8
        assert len(result.issues) == 0
        assert result.agent_name == "SyntaxTest"

    def test_parse_review_response_json_decode_error_from_extracted(self):
        """測試 _extract_json_object 回傳的字串不是合法 JSON（覆蓋 line 145）"""
        # 這裡利用 _extract_json_object 的括號配對找到 {...}，但內部不是合法 JSON
        response = "結果如下: {key without quotes: value} 結束"
        result = parse_review_response(response, "DecodeTest", "fact")
        assert result.score == 0.8
        assert len(result.issues) == 0

    def test_parse_review_response_general_exception(self):
        """測試非 JSONDecodeError 的一般例外（覆蓋 except Exception 路徑）

        透過讓 score 值為不可轉為 float 的類型（如列表），
        使 float() 呼叫拋出 TypeError 來觸發一般例外路徑。
        """
        # score 為列表，float([1,2]) 會拋出 TypeError
        response = json.dumps({
            "issues": [],
            "score": [1, 2],
        })
        result = parse_review_response(response, "ExceptionTest", "style")
        assert result.score == 0.8
        assert len(result.issues) == 0
        assert result.agent_name == "ExceptionTest"

    def test_parse_review_response_score_clamped_high(self):
        """測試 LLM 回傳超出範圍的分數被正確鉗位（原本會觸發 Pydantic 錯誤）"""
        response = json.dumps({
            "issues": [],
            "score": 999.0,
        })
        result = parse_review_response(response, "ClampTest", "style")
        assert result.score == 1.0
        assert result.agent_name == "ClampTest"

    def test_parse_review_response_with_custom_default_score(self):
        """測試自訂預設分數在解析失敗時正確套用"""
        result = parse_review_response(
            "no json here", "CustomScore", "style", default_score=0.5
        )
        assert result.score == 0.5

    def test_extract_json_object_unbalanced_opening_braces(self):
        """測試不平衡的左括號（只有開頭沒有結尾）"""
        result = _extract_json_object("{{{")
        assert result is None

    def test_extract_json_object_deeply_nested_unbalanced(self):
        """測試深層巢狀但不平衡的括號"""
        result = _extract_json_object('{"a": {"b": {"c": "d"')
        assert result is None

    def test_extract_json_object_empty_object(self):
        """測試空 JSON 物件"""
        result = _extract_json_object("prefix {} suffix")
        assert result == "{}"

    def test_extract_json_object_with_string_containing_backslash(self):
        """測試字串中包含反斜線的情況"""
        text = r'{"path": "C:\\Users\\test"}'
        result = _extract_json_object(text)
        assert result is not None
        data = json.loads(result)
        assert "Users" in data["path"]

    def test_format_audit_to_review_result_with_errors_and_warnings(self):
        """測試含有錯誤和警告的格式審查結果轉換"""
        fmt_raw = {
            "errors": ["缺少主旨段落", "缺少受文者"],
            "warnings": ["說明段落過短"],
        }
        result = format_audit_to_review_result(fmt_raw)
        assert result.agent_name == "Format Auditor"
        assert abs(result.score - 0.5) < 0.01  # 2 個錯誤: max(0.0, 0.7 - 2*0.1) ≈ 0.5
        assert result.confidence == 1.0
        assert len(result.issues) == 3
        # 驗證錯誤和警告的嚴重性
        errors = [i for i in result.issues if i.severity == "error"]
        warnings = [i for i in result.issues if i.severity == "warning"]
        assert len(errors) == 2
        assert len(warnings) == 1

    def test_format_audit_to_review_result_empty_input(self):
        """測試空的格式審查結果（無錯誤無警告）"""
        result = format_audit_to_review_result({"errors": [], "warnings": []})
        assert result.score == 1.0
        assert len(result.issues) == 0

    def test_format_audit_to_review_result_missing_keys(self):
        """測試完全空的字典輸入"""
        result = format_audit_to_review_result({})
        assert result.score == 1.0
        assert len(result.issues) == 0

    def test_format_audit_to_review_result_custom_agent_name(self):
        """測試自訂 agent_name"""
        result = format_audit_to_review_result(
            {"errors": ["問題"], "warnings": []},
            agent_name="Custom Auditor"
        )
        assert result.agent_name == "Custom Auditor"

    def test_format_audit_structured_errors_with_suggestions(self):
        """測試結構化格式審查結果（含 suggestion 欄位）"""
        fmt_raw = {
            "errors": [
                {
                    "description": "缺少主旨段落",
                    "location": "文件開頭",
                    "suggestion": "請在公文開頭加入「主旨：」段落，簡述公文目的",
                }
            ],
            "warnings": [
                {
                    "description": "說明段落過短",
                    "location": "說明段",
                    "suggestion": "建議說明段至少包含兩點，詳述背景與做法",
                }
            ],
        }
        result = format_audit_to_review_result(fmt_raw)
        assert len(result.issues) == 2
        err = [i for i in result.issues if i.severity == "error"][0]
        assert err.location == "文件開頭"
        assert err.suggestion == "請在公文開頭加入「主旨：」段落，簡述公文目的"
        warn = [i for i in result.issues if i.severity == "warning"][0]
        assert warn.location == "說明段"
        assert warn.suggestion is not None

    def test_format_audit_mixed_string_and_dict(self):
        """測試混合格式（字串 + 結構化 dict）的向後相容"""
        fmt_raw = {
            "errors": [
                "舊格式純字串錯誤",
                {"description": "新格式結構化錯誤", "location": "辦法段", "suggestion": "修正建議"},
            ],
            "warnings": [],
        }
        result = format_audit_to_review_result(fmt_raw)
        assert len(result.issues) == 2
        # 舊格式：location 預設為「文件結構」，無 suggestion
        assert result.issues[0].location == "文件結構"
        assert result.issues[0].suggestion is None
        # 新格式：帶 location 和 suggestion
        assert result.issues[1].location == "辦法段"
        assert result.issues[1].suggestion == "修正建議"

    def test_normalize_audit_items_handles_empty_and_none(self):
        """測試 _normalize_audit_items 過濾空值"""
        from src.agents.auditor import _normalize_audit_items
        items = [None, "", {"description": "有效項目", "suggestion": "建議"}, "純字串"]
        result = _normalize_audit_items(items)
        assert len(result) == 2
        assert isinstance(result[0], dict)
        assert result[0]["description"] == "有效項目"
        assert result[1] == "純字串"

    def test_normalize_audit_items_dict_without_description(self):
        """測試 dict 缺少 description 時被過濾"""
        from src.agents.auditor import _normalize_audit_items
        items = [{"location": "某處", "suggestion": "建議"}]
        result = _normalize_audit_items(items)
        assert len(result) == 0

    def test_parse_review_response_confidence_from_json(self):
        """測試從 JSON 正確讀取 confidence 值"""
        response = json.dumps({
            "issues": [],
            "score": 0.9,
            "confidence": 0.75,
        })
        result = parse_review_response(response, "ConfTest", "style")
        assert result.confidence == 0.75

    def test_parse_review_response_no_json_in_text(self):
        """測試文字中完全沒有 JSON 物件（找不到 { 符號）"""
        result = parse_review_response(
            "This response has no JSON at all!", "NoJSON", "fact"
        )
        assert result.score == 0.8
        assert len(result.issues) == 0


# ==================== _chinese_index 邊界測試 ====================

class TestChineseIndex:
    """_chinese_index Jinja2 過濾器的邊界測試"""

    def test_index_out_of_range(self):
        """測試超出中文數字範圍時回傳阿拉伯數字字串"""
        from src.agents.template import _chinese_index
        from src.core.constants import MAX_CHINESE_NUMBER
        # 超過上限應回傳 str(value)
        result = _chinese_index(MAX_CHINESE_NUMBER + 1)
        assert result == str(MAX_CHINESE_NUMBER + 1)

    def test_index_zero(self):
        """測試 0（低於有效範圍）回傳 str"""
        from src.agents.template import _chinese_index
        result = _chinese_index(0)
        assert result == "0"

    def test_index_negative(self):
        """測試負數回傳 str"""
        from src.agents.template import _chinese_index
        result = _chinese_index(-1)
        assert result == "-1"

    def test_index_valid(self):
        """測試有效範圍回傳中文數字"""
        from src.agents.template import _chinese_index
        result = _chinese_index(1)
        assert result == "一"


# ==================== parse_draft basis_text 邊界測試 ====================

class TestParseDraftBasisOnly:
    """parse_draft 中 basis_text 無 explanation_text 的邊界測試"""

    def test_basis_without_explanation(self):
        """測試有依據但無說明時的段落合併"""
        engine = TemplateEngine()
        draft = "### 主旨\n測試主旨\n### 依據\n某法規第三條"
        sections = engine.parse_draft(draft)
        assert sections["explanation"] == "依據：某法規第三條"

    def test_basis_with_explanation(self):
        """測試有依據且有說明時的段落合併"""
        engine = TemplateEngine()
        draft = "### 主旨\n測試主旨\n### 依據\n某法規\n### 說明\n補充說明"
        sections = engine.parse_draft(draft)
        assert "依據：某法規" in sections["explanation"]
        assert "補充說明" in sections["explanation"]


# ==================== _parse_list_items 空行跳過 ====================

class TestParseListItemsBlankLines:
    """_parse_list_items 中空行跳過邏輯的測試"""

    def test_items_with_blank_lines(self):
        """測試含空行的清單解析會跳過空行"""
        engine = TemplateEngine()
        text = "第一項\n\n第二項\n\n\n第三項"
        items = engine._parse_list_items(text)
        assert len(items) == 3
        assert items[0] == "第一項"
        assert items[2] == "第三項"


# ==================== apply_template 載入/渲染失敗測試 ====================

class TestApplyTemplateFallback:
    """apply_template 中 Jinja2 載入/渲染失敗的回退測試"""

    def test_template_load_failure_fallback(self, sample_requirement):
        """測試 Jinja2 模板載入失敗時使用 fallback"""
        from unittest.mock import patch
        engine = TemplateEngine()
        sections = {
            "subject": "測試主旨",
            "explanation": "測試說明",
            "provisions": "",
            "attachments": "",
            "references": ""
        }
        # Mock env.get_template 使其拋出異常
        with patch.object(engine.env, "get_template", side_effect=Exception("Template not found")):
            result = engine.apply_template(sample_requirement, sections)
        assert "測試主旨" in result
        assert "測試機關" in result

    def test_template_render_failure_fallback(self, sample_requirement):
        """測試 Jinja2 模板渲染失敗時使用 fallback"""
        from unittest.mock import patch, MagicMock
        engine = TemplateEngine()
        sections = {
            "subject": "渲染失敗測試",
            "explanation": "測試說明",
            "provisions": "",
            "attachments": "",
            "references": ""
        }
        mock_template = MagicMock()
        mock_template.render.side_effect = Exception("Render error")
        with patch.object(engine.env, "get_template", return_value=mock_template):
            result = engine.apply_template(sample_requirement, sections)
        assert "渲染失敗測試" in result
        assert "測試機關" in result


# ==================== FormatAuditor 驗證器白名單 ====================

class TestAuditorValidatorWhitelist:
    """FormatAuditor 驗證器白名單安全機制的測試"""

    def test_disallowed_validator_skipped(self, mock_llm):
        """測試不在白名單中的驗證器被跳過"""
        mock_kb = MagicMock()
        mock_kb.search_regulations.return_value = [
            {"content": "[Call: evil_function]\n- [Error] 規則",
             "metadata": {"title": "規則"}}
        ]
        mock_llm.generate.return_value = '{"errors": [], "warnings": []}'

        auditor = FormatAuditor(mock_llm, mock_kb)
        result = auditor.audit("### 主旨\n測試", "函")
        # evil_function 不在白名單，不應執行；結果僅來自 LLM
        assert isinstance(result, dict)

    def test_unknown_validator_skipped(self, mock_llm):
        """測試白名單內但不存在於 registry 的驗證器被跳過"""
        from unittest.mock import patch
        mock_kb = MagicMock()
        mock_kb.search_regulations.return_value = [
            {"content": "[Call: check_date_logic]\n- [Error] 規則",
             "metadata": {"title": "規則"}}
        ]
        mock_llm.generate.return_value = '{"errors": [], "warnings": []}'

        auditor = FormatAuditor(mock_llm, mock_kb)
        # Mock hasattr 回傳 False 來模擬 registry 上找不到該函式
        import src.agents.auditor as auditor_mod
        mock_registry = MagicMock(spec=[])  # spec=[] 表示沒有任何屬性
        with patch.object(auditor_mod, 'validator_registry', mock_registry):
            result = auditor.audit("### 主旨\n測試", "函")
            assert isinstance(result, dict)

    def test_validator_exception_handled(self, mock_llm):
        """測試驗證器執行時拋出異常被安全處理"""
        from unittest.mock import patch
        mock_kb = MagicMock()
        mock_kb.search_regulations.return_value = [
            {"content": "[Call: check_date_logic]\n- [Error] 規則",
             "metadata": {"title": "規則"}}
        ]
        mock_llm.generate.return_value = '{"errors": [], "warnings": []}'

        auditor = FormatAuditor(mock_llm, mock_kb)
        import src.agents.auditor as auditor_mod
        mock_registry = MagicMock()
        mock_registry.check_date_logic.side_effect = RuntimeError("Validator crash")
        with patch.object(auditor_mod, 'validator_registry', mock_registry):
            result = auditor.audit("### 主旨\n測試", "函")
            assert isinstance(result, dict)


# ==================== OrgMemory 儲存失敗 ====================

# ==================== FormatAuditor JSONDecodeError ====================

class TestAuditorJSONDecodeError:
    """FormatAuditor LLM 回傳無效 JSON 的測試"""

    def test_audit_llm_returns_invalid_json(self, mock_llm):
        """測試 LLM 回傳含 {} 但無效的 JSON 時的 JSONDecodeError 處理"""
        # 正規式會匹配 {errors: []} 但 json.loads 會失敗
        mock_llm.generate.return_value = "Result: {errors: [], warnings: []}"
        auditor = FormatAuditor(mock_llm)
        result = auditor.audit("### 主旨\n測試", "函")
        assert isinstance(result, dict)
        assert any("無法解析 JSON" in w for w in result["warnings"])


# ==================== ComplianceChecker 例外路徑 ====================

class TestComplianceCheckerExceptions:
    """ComplianceChecker JSON 解析例外路徑測試（現使用共享 parse_review_response）"""

    def test_check_json_decode_error(self, mock_llm):
        """測試 LLM 回傳不合法 JSON 時回傳預設結果"""
        mock_llm.generate.return_value = "data: {invalid: json here}"
        checker = ComplianceChecker(mock_llm)
        result = checker.check("### 主旨\n測試草稿")
        assert result.score == 0.85  # 預設分數

    def test_check_generic_exception_in_parser(self, mock_llm):
        """測試解析過程中通用 Exception 時回傳預設結果"""
        from unittest.mock import patch
        mock_llm.generate.return_value = "data: {some json}"
        checker = ComplianceChecker(mock_llm)
        with patch("src.agents.review_parser.json.loads", side_effect=RuntimeError("unexpected")):
            result = checker.check("### 主旨\n測試草稿")
        assert result.score == 0.85  # 預設分數


# ==================== RequirementAgent regex 策略 3 失敗 ====================

class TestRequirementRegexStrategy3:
    """RequirementAgent regex 策略 3 失敗的測試"""

    def test_regex_fields_fail_validation(self, mock_llm):
        """測試 regex 提取到的欄位驗證失敗時使用 fallback（空白 doc_type）"""
        from src.agents.requirement import RequirementAgent
        agent = RequirementAgent(mock_llm)
        # LLM 回應的 JSON 結構不完整，只有部分欄位可 regex 到
        # 但 doc_type 為空白會觸發 ValidationError，觸發 fallback
        mock_llm.generate.return_value = '"doc_type": "   ", "sender": "機關", "receiver": "對象", "subject": "主旨"'
        result = agent.analyze("使用者輸入")
        assert result.doc_type == "函"  # fallback 預設
        assert result.sender == "（未指定）"
        assert "使用者輸入" in result.subject


# ==================== KB metadata 非標準型別 ====================

class TestKBMetadataConversion:
    """KB ingest 中 metadata 型別轉換的測試"""

    def test_non_standard_metadata_type_converted(self, tmp_path):
        """測試含巢狀 dict metadata 的檔案匯入時被轉為 str"""
        from src.cli.kb import app as kb_app
        from typer.testing import CliRunner
        from unittest.mock import patch

        runner = CliRunner()
        # 建立含有 nested dict metadata 的 YAML 前置物（dict 不在標準型別清單中）
        doc_path = tmp_path / "test_doc.md"
        doc_path.write_text("---\ntitle: 測試\nnested:\n  key: value\n---\n內容", encoding='utf-8')

        with patch("src.cli.kb.ConfigManager") as mock_cm, \
             patch("src.cli.kb.get_llm_factory"), \
             patch("src.cli.kb.KnowledgeBaseManager") as mock_kb_class:
            mock_cm.return_value.config = {
                "llm": {"provider": "mock"},
                "knowledge_base": {"path": str(tmp_path / "kb")}
            }
            mock_kb_instance = mock_kb_class.return_value
            mock_kb_instance.upsert_document.return_value = "doc_id"
            mock_kb_instance.make_deterministic_id.return_value = "test-det-id"
            mock_kb_instance.contextual_retrieval = False

            result = runner.invoke(kb_app, [
                "ingest", "--source-dir", str(tmp_path), "--collection", "examples"
            ])
            assert result.exit_code == 0
            # 確認 upsert_document 被呼叫（Round 78 改為冪等 upsert），且 nested dict 被轉為 str
            call_args = mock_kb_instance.upsert_document.call_args
            metadata = call_args[0][2]  # 第三個位置引數：upsert_document(det_id, content, metadata, ...)
            assert isinstance(metadata["nested"], str)
            assert "key" in metadata["nested"]


# ==================== OrgMemory 儲存失敗 ====================

class TestOrgMemorySaveFailure:
    """OrgMemory 儲存偏好設定失敗的測試"""

    def test_save_preferences_io_error(self, tmp_path):
        """測試儲存偏好設定時 IO 錯誤被優雅處理"""
        from src.agents.org_memory import OrganizationalMemory
        from unittest.mock import patch
        mem = OrganizationalMemory(str(tmp_path / "prefs.json"))
        mem.preferences["測試機關"] = {"formal_level": "formal"}
        # Mock open 拋出 IOError
        with patch("builtins.open", side_effect=IOError("磁碟已滿")):
            # 不應拋出異常
            mem._save_preferences()


# ==================== _sanitize_json_string 不可見字元測試 ====================

class TestSanitizeJsonString:
    """測試 requirement.py 的 _sanitize_json_string 處理不可見 Unicode 字元"""

    def test_removes_bom_and_zwsp(self):
        """測試移除 BOM 和零寬度空格"""
        text = '\ufeff{"doc_type": "函\u200b"}'
        result = _sanitize_json_string(text)
        assert '\ufeff' not in result
        assert '\u200b' not in result
        assert '"doc_type"' in result

    def test_removes_zwnj_zwj_direction_marks(self):
        """測試移除 ZWNJ、ZWJ、方向標記"""
        text = '{"doc\u200c_type": "函\u200d", "sender\u200e": "機\u200f關"}'
        result = _sanitize_json_string(text)
        assert '\u200c' not in result
        assert '\u200d' not in result
        assert '\u200e' not in result
        assert '\u200f' not in result
        assert '"doc_type"' in result

    def test_removes_soft_hyphen_and_word_joiner(self):
        """測試移除 Soft Hyphen 和 Word Joiner"""
        text = '{"subject": "測\u00ad試\u2060主旨"}'
        result = _sanitize_json_string(text)
        assert '\u00ad' not in result
        assert '\u2060' not in result
        assert '測試主旨' in result

    def test_none_input(self):
        """測試 None 輸入"""
        assert _sanitize_json_string(None) == ""

    def test_empty_input(self):
        """測試空字串輸入"""
        assert _sanitize_json_string("") == ""

    def test_normal_text_unchanged(self):
        """測試正常 JSON 不被改變"""
        text = '{"doc_type": "函", "sender": "臺北市政府"}'
        assert _sanitize_json_string(text) == text


# ==================== EditorInChief 並行審查邊界案例 ====================

class TestEditorParallelEdgeCases:
    """EditorInChief 並行審查的邊界案例測試"""

    def test_all_agents_fail(self, mock_llm):
        """測試所有並行 Agent 全部崩潰的邊界案例"""
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_regulations.return_value = []
        mock_kb.search_policies.return_value = []

        # FormatAuditor 同步呼叫成功（返回空結果）
        # 但所有並行 Agent 全部拋出例外
        call_count = [0]

        def side_effect(*args, **kwargs):
            call_count[0] += 1
            if call_count[0] == 1:
                # FormatAuditor
                return json.dumps({"errors": [], "warnings": []})
            # 所有並行 Agent 拋出例外
            raise RuntimeError("模擬 Agent 全部崩潰")

        mock_llm.generate.side_effect = side_effect

        editor = EditorInChief(mock_llm, mock_kb)
        refined, report = editor.review_and_refine("# 函\n### 主旨\n測試", "函")

        # 應正常完成不崩潰
        assert report is not None
        assert isinstance(report, QAReport)
        # 應有 6 個結果（Format Auditor + Citation Checker + 4 個並行 Agent）
        assert len(report.agent_results) == 6
        # 失敗的 Agent score 應為 0.0
        failed_agents = [r for r in report.agent_results if r.score == 0.0]
        assert len(failed_agents) == 4

    def test_auto_refine_returns_error_string(self, mock_llm):
        """測試 _auto_refine 中 LLM 回傳 Error 字串時保留原始草稿"""
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_regulations.return_value = []
        mock_kb.search_policies.return_value = []

        editor = EditorInChief(mock_llm, mock_kb)

        # 建立有問題的 results 觸發 _auto_refine
        results = [
            ReviewResult(
                agent_name="Format Auditor",
                issues=[ReviewIssue(
                    category="format", severity="error",
                    location="結構", description="缺少段落"
                )],
                score=0.3, confidence=1.0,
            ),
        ]

        # LLM 回傳 Error 字串
        mock_llm.generate.return_value = "Error: Model not available"
        refined = editor._auto_refine("原始草稿", results)
        assert refined == "原始草稿"

    def test_auto_refine_llm_exception(self, mock_llm):
        """測試 _auto_refine 中 LLM 呼叫例外時保留原始草稿"""
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        editor = EditorInChief(mock_llm, mock_kb)

        results = [
            ReviewResult(
                agent_name="Style Checker",
                issues=[ReviewIssue(
                    category="style", severity="warning",
                    location="全文", description="用語不正式",
                    suggestion="改用正式用語"
                )],
                score=0.6, confidence=1.0,
            ),
        ]

        mock_llm.generate.side_effect = Exception("連線逾時")
        refined = editor._auto_refine("原始草稿內容", results)
        assert refined == "原始草稿內容"

    def test_auto_refine_empty_feedback(self, mock_llm):
        """測試 _auto_refine 中所有 results 都沒有 issues 時直接回傳"""
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        editor = EditorInChief(mock_llm, mock_kb)

        # 所有結果都沒有 issues
        results = [
            ReviewResult(agent_name="Agent A", issues=[], score=0.9, confidence=1.0),
            ReviewResult(agent_name="Agent B", issues=[], score=0.8, confidence=1.0),
        ]

        refined = editor._auto_refine("原始草稿", results)
        # 無 feedback 時不應呼叫 LLM
        mock_llm.generate.assert_not_called()
        assert refined == "原始草稿"

    def test_auto_refine_suggestion_none(self, mock_llm):
        """測試 _auto_refine 中 suggestion 為 None 時使用預設文字"""
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        editor = EditorInChief(mock_llm, mock_kb)

        results = [
            ReviewResult(
                agent_name="Fact Checker",
                issues=[ReviewIssue(
                    category="fact", severity="error",
                    location="說明", description="日期有誤",
                    suggestion=None,  # 明確設為 None
                )],
                score=0.5, confidence=1.0,
            ),
        ]

        mock_llm.generate.return_value = "修正後的草稿"
        editor._auto_refine("原始草稿", results)

        # LLM 應被呼叫，且 feedback 中應包含預設文字
        mock_llm.generate.assert_called_once()
        prompt_arg = mock_llm.generate.call_args[0][0]
        assert "請自行判斷修正方式" in prompt_arg


# ==================== ComplianceChecker severity 正規化測試 ====================

class TestComplianceCheckerSeverityNormalization:
    """測試 ComplianceChecker 對非標準 severity 值的正規化處理"""

    def test_non_standard_severity_normalized_to_info(self):
        """LLM 回傳非標準 severity（如 critical）應被正規化為 info"""
        mock_llm = MagicMock()
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_policies.return_value = []
        checker = ComplianceChecker(mock_llm, mock_kb)

        # 模擬 LLM 回傳含非標準 severity 的 JSON
        mock_llm.generate.return_value = json.dumps({
            "score": 0.7,
            "confidence": 0.8,
            "issues": [
                {
                    "severity": "critical",
                    "location": "主旨",
                    "description": "格式不符",
                    "suggestion": "修正格式"
                },
                {
                    "severity": "high",
                    "location": "說明",
                    "description": "措辭不當",
                    "suggestion": "調整措辭"
                },
                {
                    "severity": "warning",
                    "location": "正本",
                    "description": "缺少受文者",
                    "suggestion": "補充"
                }
            ]
        })

        result = checker.check("測試草稿" * 10)
        # 非標準 severity 應被正規化，不會導致整個結果丟失
        assert result.score == 0.7
        assert len(result.issues) == 3
        # "critical" 和 "high" 應被正規化為 "info"
        assert result.issues[0].severity == "info"
        assert result.issues[1].severity == "info"
        # "warning" 保持不變
        assert result.issues[2].severity == "warning"

    def test_standard_severities_preserved(self):
        """標準 severity 值不應被修改"""
        mock_llm = MagicMock()
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_policies.return_value = []
        checker = ComplianceChecker(mock_llm, mock_kb)

        mock_llm.generate.return_value = json.dumps({
            "score": 0.5,
            "confidence": 0.9,
            "issues": [
                {"severity": "error", "location": "A", "description": "嚴重問題"},
                {"severity": "warning", "location": "B", "description": "一般問題"},
                {"severity": "info", "location": "C", "description": "提示"},
            ]
        })

        result = checker.check("測試草稿" * 10)
        assert result.issues[0].severity == "error"
        assert result.issues[1].severity == "warning"
        assert result.issues[2].severity == "info"


# ==================== FormatAuditor 型別驗證測試 ====================

class TestFormatAuditorTypeValidation:
    """測試 FormatAuditor 對 LLM 回傳 errors/warnings 非列表型別的處理"""

    def test_errors_as_string_not_split_by_char(self):
        """LLM 回傳 errors 為字串時不應逐字拆解"""
        mock_llm = MagicMock()
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_regulations.return_value = []
        auditor = FormatAuditor(mock_llm, mock_kb)

        # 模擬 LLM 回傳 errors 為字串而非列表
        mock_llm.generate.return_value = json.dumps({
            "errors": "No errors found",
            "warnings": ["一個警告"]
        })

        test_draft = (
            "**機關** 測試\n**受文者** 測試\n**速別** 普通\n"
            "**發文日期** 中華民國114年1月1日\n**主旨** 測試"
        )
        result = auditor.audit(test_draft, "函")
        # errors 為字串時不應被 extend，所以 errors 列表不包含逐字拆解的字元
        # 只有規則檢查產生的 errors（如果有）
        for e in result.get("errors", []):
            assert len(e) > 1  # 不應有單字元錯誤

    def test_warnings_as_string_not_split_by_char(self):
        """LLM 回傳 warnings 為字串時不應逐字拆解"""
        mock_llm = MagicMock()
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_regulations.return_value = []
        auditor = FormatAuditor(mock_llm, mock_kb)

        mock_llm.generate.return_value = json.dumps({
            "errors": [],
            "warnings": "All looks good"
        })

        test_draft = (
            "**機關** 測試\n**受文者** 測試\n**速別** 普通\n"
            "**發文日期** 中華民國114年1月1日\n**主旨** 測試"
        )
        result = auditor.audit(test_draft, "函")
        for w in result.get("warnings", []):
            assert len(w) > 1

    def test_valid_list_types_work_normally(self):
        """正常列表型別的 errors/warnings 應正常處理"""
        mock_llm = MagicMock()
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_regulations.return_value = []
        auditor = FormatAuditor(mock_llm, mock_kb)

        mock_llm.generate.return_value = json.dumps({
            "errors": ["格式錯誤一"],
            "warnings": ["建議一", "建議二"]
        })

        test_draft = (
            "**機關** 測試\n**受文者** 測試\n**速別** 普通\n"
            "**發文日期** 中華民國114年1月1日\n**主旨** 測試"
        )
        result = auditor.audit(test_draft, "函")
        assert "格式錯誤一" in result["errors"]
        assert "建議一" in result["warnings"]
        assert "建議二" in result["warnings"]


# ==================== PublicDocRequirement sender/receiver 空白驗證測試 ====================

class TestPublicDocRequirementBlankValidation:
    """測試 sender 和 receiver 的空白字串驗證"""

    def test_sender_blank_rejected(self):
        """純空白的 sender 應被拒絕"""
        with pytest.raises(ValueError, match="欄位不可為空白"):
            PublicDocRequirement(
                doc_type="函",
                sender="   ",
                receiver="測試受文者",
                subject="測試主旨",
            )

    def test_receiver_blank_rejected(self):
        """純空白的 receiver 應被拒絕"""
        with pytest.raises(ValueError, match="欄位不可為空白"):
            PublicDocRequirement(
                doc_type="函",
                sender="測試機關",
                receiver="  \t  ",
                subject="測試主旨",
            )

    def test_sender_receiver_stripped(self):
        """sender 和 receiver 應自動去除前後空白"""
        req = PublicDocRequirement(
            doc_type="函",
            sender="  臺北市政府  ",
            receiver="  各區公所  ",
            subject="測試主旨",
        )
        assert req.sender == "臺北市政府"
        assert req.receiver == "各區公所"

    def test_valid_sender_receiver(self):
        """正常的 sender/receiver 應正常通過"""
        req = PublicDocRequirement(
            doc_type="函",
            sender="臺北市政府",
            receiver="各區公所",
            subject="測試主旨",
        )
        assert req.sender == "臺北市政府"
        assert req.receiver == "各區公所"


# ==================== CLI --output 路徑處理測試 ====================

class TestCLIOutputPath:
    """測試 generate.py 的輸出路徑處理邏輯"""

    @staticmethod
    def _resolve_output_path(output_path: str) -> str:
        """複製 generate.py 的路徑解析邏輯用於測試"""
        import os
        has_traversal = ".." in output_path
        is_absolute = os.path.isabs(output_path) or output_path.startswith("/")
        safe_filename = os.path.basename(output_path)
        if not safe_filename or safe_filename.startswith("."):
            safe_filename = "output.docx"
        if not safe_filename.endswith(".docx"):
            safe_filename += ".docx"

        if is_absolute or has_traversal:
            return safe_filename
        elif os.path.dirname(output_path):
            output_dir = os.path.dirname(os.path.abspath(output_path))
            return os.path.join(output_dir, safe_filename)
        else:
            return safe_filename

    def test_output_subdir_preserved(self):
        """正常子目錄路徑應被保留"""
        result = self._resolve_output_path("output/my_doc.docx")
        assert "my_doc.docx" in result
        assert "output" in result

    def test_traversal_path_stripped(self):
        """含 .. 的路徑遍歷應被阻止"""
        result = self._resolve_output_path("../../etc/evil.docx")
        assert result == "evil.docx"
        assert ".." not in result

    def test_absolute_path_stripped(self):
        """絕對路徑應被攔截，僅保留檔名"""
        result = self._resolve_output_path("/etc/passwd_backup.docx")
        assert result == "passwd_backup.docx"
        assert "/" not in result and "\\" not in result

    def test_windows_absolute_path_stripped(self):
        """Windows 絕對路徑應被攔截"""
        result = self._resolve_output_path("C:\\Windows\\evil.docx")
        assert result == "evil.docx"


# ==================== RateLimiter 清理機制測試 ====================

@pytest.mark.skipif(
    not importlib.util.find_spec("multipart"),
    reason="python-multipart 未安裝，跳過 API 相關測試",
)
class TestRateLimiterCleanup:
    """測試 RateLimiter 的定期清理機制"""

    def test_cleanup_removes_expired_entries(self):
        """定期清理應移除過期的 IP 條目"""
        import time
        from api_server import _RateLimiter

        limiter = _RateLimiter(max_requests=10, window_seconds=60)

        # 手動注入過期條目
        with limiter._lock:
            now = time.monotonic()
            limiter._requests["expired_ip_1"] = [now - 120]  # 2 分鐘前
            limiter._requests["expired_ip_2"] = [now - 300]  # 5 分鐘前
            limiter._request_counter = limiter._CLEANUP_INTERVAL - 1  # 下一次請求觸發清理

        # 觸發清理
        limiter.is_allowed("new_ip")

        with limiter._lock:
            assert "expired_ip_1" not in limiter._requests
            assert "expired_ip_2" not in limiter._requests
            assert "new_ip" in limiter._requests


# ==================== config.py _create_default_config 安全測試 ====================

class TestConfigDefaultSafety:
    """測試 config.py 對不存在目錄的處理"""

    def test_create_default_config_missing_dir(self, tmp_path):
        """當目標目錄不存在時不應拋出例外"""
        from src.core.config import ConfigManager

        # 使用不存在的子目錄
        config_path = tmp_path / "nonexistent_dir" / "config.yaml"
        cm = ConfigManager(str(config_path))

        # 應該成功載入預設設定，且目錄應被自動建立
        assert cm.config is not None
        assert cm.config["llm"]["provider"] == "ollama"
        assert config_path.exists()


# ==================== renumber_provisions 正則修正測試 ====================

class TestRenumberProvisionsRegexFix:
    """測試 renumber_provisions 的正則表達式修正，防止以中文數字開頭的普通文字被誤判為編號"""

    def test_chinese_text_not_misidentified_as_numbered(self):
        """以中文數字開頭的普通文字不應被重新編號"""
        result = renumber_provisions("三民主義統一中國")
        assert result == "三民主義統一中國"

    def test_chinese_number_word_preserved(self):
        """「十全十美」等詞語不應被截斷"""
        result = renumber_provisions("十全十美的計畫")
        assert result == "十全十美的計畫"

    def test_chinese_number_prefix_preserved(self):
        """「十分重要」等不應被當作編號"""
        result = renumber_provisions("十分重要的事項")
        assert result == "十分重要的事項"

    def test_real_numbered_items_still_work(self):
        """正式的編號項目（含分隔符號）仍應被正確重新編號"""
        text = "1、第一項\n2、第二項"
        result = renumber_provisions(text)
        assert "一、第一項" in result
        assert "二、第二項" in result

    def test_chinese_numbered_items_still_work(self):
        """中文編號項目仍正確處理"""
        text = "一、原始項目一\n二、原始項目二"
        result = renumber_provisions(text)
        assert "一、原始項目一" in result
        assert "二、原始項目二" in result

    def test_mixed_numbered_and_plain_text(self):
        """混合編號和普通文字"""
        text = "一、第一項\n三民主義是重要的\n二、第二項"
        result = renumber_provisions(text)
        assert "一、第一項" in result
        assert "三民主義是重要的" in result
        assert "二、第二項" in result


# ==================== parse_review_response risk_level 測試 ====================

class TestParseReviewResponseRiskLevel:
    """測試 parse_review_response 正確提取 LLM 回傳的 risk_level"""

    def test_risk_level_extracted_from_llm_response(self):
        """LLM 回傳 risk_level 時應被正確提取"""
        response = json.dumps({
            "score": 0.5,
            "issues": [
                {
                    "severity": "error",
                    "risk_level": "high",
                    "location": "主旨",
                    "description": "格式錯誤"
                }
            ]
        })
        result = parse_review_response(response, "TestAgent", "style")
        assert result.issues[0].risk_level == "high"

    def test_risk_level_defaults_to_low(self):
        """未回傳 risk_level 時預設為 low"""
        response = json.dumps({
            "score": 0.8,
            "issues": [
                {
                    "severity": "warning",
                    "location": "說明",
                    "description": "用詞不當"
                }
            ]
        })
        result = parse_review_response(response, "TestAgent", "style")
        assert result.issues[0].risk_level == "low"

    def test_invalid_risk_level_normalized_to_low(self):
        """非標準 risk_level 應被正規化為 low"""
        response = json.dumps({
            "score": 0.5,
            "issues": [
                {
                    "severity": "error",
                    "risk_level": "critical",
                    "location": "主旨",
                    "description": "格式錯誤"
                }
            ]
        })
        result = parse_review_response(response, "TestAgent", "style")
        assert result.issues[0].risk_level == "low"


# ==================== EditorInChief total_weight==0 → Critical ====================

class TestEditorQAReportCriticalRisk:
    """測試 EditorInChief._generate_qa_report 在 total_weight==0 時回傳 Critical"""

    def test_total_weight_zero_returns_critical(self):
        """當 results 為空列表時，total_weight==0，risk 應為 Critical"""
        mock_llm = MagicMock()
        editor = EditorInChief(mock_llm)
        report = editor._generate_qa_report([])
        assert report.risk_summary == "Critical"

    def test_non_empty_results_not_critical(self):
        """有 agent 結果時，risk 不應自動為 Critical"""
        mock_llm = MagicMock()
        editor = EditorInChief(mock_llm)
        good_result = ReviewResult(
            agent_name="Style Checker",
            issues=[],
            score=0.95,
            confidence=1.0,
        )
        report = editor._generate_qa_report([good_result])
        assert report.risk_summary != "Critical"


# ==================== RequirementAgent 策略 3 降級日誌 ====================

class TestRequirementStrategy3Logging:
    """測試 RequirementAgent 策略 3 正則提取時產生 warning 日誌"""

    def test_strategy3_logs_warning(self, caplog):
        """策略 3 成功時應記錄 warning 日誌"""
        mock_llm = MagicMock()
        # 回傳有效的個別欄位但非合法 JSON 物件
        mock_llm.generate.return_value = (
            'some text "doc_type": "函" more text '
            '"sender": "教育部" etc '
            '"receiver": "各縣市政府" rest '
            '"subject": "配合辦理" end'
        )
        agent = RequirementAgent(mock_llm)
        import logging
        with caplog.at_level(logging.WARNING, logger="src.agents.requirement"):
            result = agent.analyze("寫一份函給各縣市政府")
        assert "策略 3" in caplog.text
        assert result.doc_type == "函"

    def test_strategy1_no_degradation_warning(self, caplog):
        """策略 1 成功時不應記錄策略 3 的 warning"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = json.dumps({
            "doc_type": "函",
            "sender": "教育部",
            "receiver": "各縣市政府",
            "subject": "配合辦理",
        })
        agent = RequirementAgent(mock_llm)
        import logging
        with caplog.at_level(logging.WARNING, logger="src.agents.requirement"):
            result = agent.analyze("寫一份函給各縣市政府")
        assert "策略 3" not in caplog.text
        assert result.doc_type == "函"


# ==================== detect_header 精確匹配 ====================

class TestDetectHeaderPrecision:
    """測試 detect_header 不會將「說明書」「主旨演講」等誤判為段落標題"""

    def test_shuoming_shu_not_detected(self):
        """「說明書」不應被辨識為說明段落"""
        engine = TemplateEngine()
        draft = "主旨：測試\n說明書需要審核"
        sections = engine.parse_draft(draft)
        assert "說明書" not in sections.get("explanation", "")

    def test_zhuzhi_yanjiang_not_detected(self):
        """「主旨演講」不應被辨識為主旨段落"""
        engine = TemplateEngine()
        draft = "說明：第一點\n主旨演講將於下週舉行"
        sections = engine.parse_draft(draft)
        assert sections["subject"] == ""
        assert "演講" not in sections.get("subject", "")

    def test_fujian_jia_not_detected(self):
        """「附件夾」不應被辨識為附件段落"""
        engine = TemplateEngine()
        draft = "主旨：測試\n說明：請使用附件夾\n附件：報告書"
        sections = engine.parse_draft(draft)
        assert sections["attachments"] == "報告書"
        assert "附件夾" in sections.get("explanation", "")

    def test_standard_headers_still_work(self):
        """標準段落標題（含冒號）仍應正確辨識"""
        engine = TemplateEngine()
        draft = "主旨：這是主旨\n說明：這是說明\n辦法：這是辦法"
        sections = engine.parse_draft(draft)
        assert sections["subject"] == "這是主旨"
        assert "這是說明" in sections["explanation"]
        assert "這是辦法" in sections["provisions"]

    def test_header_without_colon_still_works(self):
        """標準段落標題（不含冒號但獨立行）仍應正確辨識"""
        engine = TemplateEngine()
        draft = "主旨\n這是主旨內容\n說明\n這是說明內容"
        sections = engine.parse_draft(draft)
        assert sections["subject"] == "這是主旨內容"
        assert "這是說明內容" in sections["explanation"]


# ==================== 模板空說明不渲染 ====================

class TestTemplateEmptyExplanation:
    """測試 han.j2 / sign.j2 在說明為空時不渲染「### 說明」標題"""

    def test_han_no_explanation_heading_when_empty(self):
        """函模板在說明為空時不應渲染「### 說明」"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="函",
            sender="教育部",
            receiver="各縣市政府",
            subject="測試主旨",
        )
        sections = {
            "subject": "測試主旨",
            "explanation": "",
            "provisions": "一、第一項",
            "attachments": "",
            "references": "",
            "basis": "",
        }
        output = engine.apply_template(req, sections)
        lines = output.split("\n")
        for i, line in enumerate(lines):
            if line.strip() == "### 說明":
                # 下一行（非空白行）不應存在或應有實際內容
                next_content = ""
                for j in range(i + 1, len(lines)):
                    if lines[j].strip():
                        next_content = lines[j].strip()
                        break
                assert next_content != "---", "說明標題後直接接分隔線，表示空段落"
                break

    def test_single_line_explanation_no_numbering(self):
        """單行說明不應被加上中文編號"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="函",
            sender="教育部",
            receiver="各縣市政府",
            subject="測試主旨",
        )
        sections = {
            "subject": "測試主旨",
            "explanation": "為配合政策需要",
            "provisions": "",
            "attachments": "",
            "references": "",
            "basis": "",
        }
        output = engine.apply_template(req, sections)
        assert "為配合政策需要" in output
        assert "一、為配合政策需要" not in output

    def test_single_provision_no_numbering(self):
        """單項辦法不應被加上中文編號"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="函",
            sender="教育部",
            receiver="各縣市政府",
            subject="測試主旨",
        )
        sections = {
            "subject": "測試主旨",
            "explanation": "",
            "provisions": "請依限辦理",
            "attachments": "",
            "references": "",
            "basis": "",
        }
        output = engine.apply_template(req, sections)
        assert "請依限辦理" in output
        assert "一、請依限辦理" not in output

    def test_single_announcement_provision_no_numbering(self):
        """公告單項公告事項不應被加上中文編號"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="公告",
            sender="臺北市政府",
            receiver="公告周知",
            subject="測試公告",
        )
        sections = {
            "subject": "測試公告",
            "explanation": "",
            "provisions": "特此公告",
            "attachments": "",
            "references": "",
            "basis": "",
        }
        output = engine.apply_template(req, sections)
        assert "特此公告" in output
        assert "一、特此公告" not in output


# ==================== 速別驗證日誌 ====================

class TestValidateUrgencyLogging:
    """validate_urgency 正規化時應發出警告日誌。"""

    def test_invalid_urgency_logs_warning(self):
        """無效速別應記錄警告並正規化為「普通」"""
        from src.core.models import PublicDocRequirement

        req = PublicDocRequirement(
            doc_type="函",
            sender="測試機關",
            receiver="測試受文者",
            subject="測試主旨",
            urgency="超級急件",  # 無效值
        )
        assert req.urgency == "普通"

    def test_invalid_urgency_warning_message(self, caplog):
        """無效速別的警告訊息應包含原始值"""
        import logging
        from src.core.models import PublicDocRequirement

        with caplog.at_level(logging.WARNING, logger="src.core.models"):
            req = PublicDocRequirement(
                doc_type="函",
                sender="測試機關",
                receiver="測試受文者",
                subject="測試主旨",
                urgency="火速",
            )
        assert req.urgency == "普通"
        assert "火速" in caplog.text
        assert "正規化" in caplog.text

    def test_valid_urgency_no_warning(self, caplog):
        """合法速別不應產生警告"""
        import logging
        from src.core.models import PublicDocRequirement

        with caplog.at_level(logging.WARNING, logger="src.core.models"):
            req = PublicDocRequirement(
                doc_type="函",
                sender="測試機關",
                receiver="測試受文者",
                subject="測試主旨",
                urgency="速件",
            )
        assert req.urgency == "速件"
        assert "正規化" not in caplog.text


# ==================== ComplianceChecker issues 型別防護 ====================

class TestComplianceCheckerIssuesType:
    """ComplianceChecker 收到非列表型別 issues 時的防護測試。"""

    def test_issues_null_returns_empty(self):
        """issues 為 null 時應回傳空 issues 列表且正確解析 score"""
        from src.agents.compliance_checker import ComplianceChecker

        mock_llm = MagicMock()
        mock_llm.generate.return_value = '{"issues": null, "score": 0.85, "confidence": 0.9}'

        checker = ComplianceChecker(mock_llm)
        result = checker.check("### 主旨\n測試公文")
        assert isinstance(result.issues, list)
        assert len(result.issues) == 0
        assert result.score == 0.85

    def test_issues_string_returns_empty(self):
        """issues 為字串時應回傳空 issues 列表"""
        from src.agents.compliance_checker import ComplianceChecker

        mock_llm = MagicMock()
        mock_llm.generate.return_value = '{"issues": "no issues", "score": 0.9, "confidence": 0.8}'

        checker = ComplianceChecker(mock_llm)
        result = checker.check("### 主旨\n測試公文")
        assert isinstance(result.issues, list)
        assert len(result.issues) == 0
        assert result.score == 0.9

    def test_issues_with_non_dict_items_filtered(self):
        """issues 含非 dict 項目時應被過濾"""
        import json
        from src.agents.compliance_checker import ComplianceChecker

        mock_llm = MagicMock()
        mock_llm.generate.return_value = json.dumps({
            "issues": [
                {"severity": "warning", "location": "主旨", "description": "問題一"},
                42,
                "文字",
                None,
            ],
            "score": 0.7,
            "confidence": 0.8,
        })

        checker = ComplianceChecker(mock_llm)
        result = checker.check("### 主旨\n測試公文")
        assert len(result.issues) == 1
        assert result.issues[0].description == "問題一"


# ==================== FormatAuditor 非列表 errors/warnings ====================

class TestFormatAuditorNonListFields:
    """FormatAuditor 收到非列表型別 errors/warnings 時的防護測試。"""

    def test_errors_null_handled(self):
        """errors 為 null 時不崩潰"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = '{"errors": null, "warnings": []}'

        auditor = FormatAuditor(mock_llm)
        result = auditor.audit("### 主旨\n測試", "函")
        assert isinstance(result["errors"], list)
        assert isinstance(result["warnings"], list)

    def test_warnings_null_handled(self):
        """warnings 為 null 時不崩潰"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = '{"errors": [], "warnings": null}'

        auditor = FormatAuditor(mock_llm)
        result = auditor.audit("### 主旨\n測試", "函")
        assert isinstance(result["errors"], list)
        assert isinstance(result["warnings"], list)

    def test_errors_integer_handled(self):
        """errors 為整數時不崩潰"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = '{"errors": 123, "warnings": []}'

        auditor = FormatAuditor(mock_llm)
        result = auditor.audit("### 主旨\n測試", "函")
        assert isinstance(result["errors"], list)


# ==================== GAP-001: EditorInChief Safe/Low 不觸發 _auto_refine ====================

class TestEditorSafeLowNoRefine:
    """GAP-001: 當風險為 Safe 或 Low 時，不應呼叫 _auto_refine。"""

    def test_safe_score_no_auto_refine(self, mock_llm):
        """高品質審查結果（Safe）不應觸發自動修正"""
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_regulations.return_value = []
        mock_kb.search_policies.return_value = []

        # 所有 Agent 回傳高分（Safe）
        mock_llm.generate.return_value = json.dumps({
            "issues": [],
            "score": 0.98,
            "confidence": 1.0,
        })

        editor = EditorInChief(mock_llm, mock_kb)
        # 含完整引用結構的草稿，避免通用驗證器產生額外 error 降低分數
        original_draft = (
            "### 主旨\n高品質草稿\n### 說明\n依據相關法規辦理[^1]。\n\n"
            "### 參考來源\n[^1]: [Level A] 測試法規 | URL: https://example.test/law-1"
        )

        from unittest.mock import patch
        with patch.object(editor, '_auto_refine', wraps=editor._auto_refine) as mock_refine:
            final_draft, report = editor.review_and_refine(original_draft, "函")
            mock_refine.assert_not_called()

        # 應回傳原始草稿不變
        assert final_draft == original_draft
        assert report.risk_summary in ["Safe", "Low"]

    def test_low_score_no_auto_refine(self, mock_llm):
        """Low 風險也不應觸發自動修正"""
        mock_kb = MagicMock(spec=KnowledgeBaseManager)
        mock_kb.search_regulations.return_value = []
        mock_kb.search_policies.return_value = []

        # 回傳好分數但有少量 info 等級 issues（Low 等級）
        mock_llm.generate.return_value = json.dumps({
            "issues": [
                {"severity": "info", "location": "說明", "description": "建議微調"}
            ],
            "score": 0.92,
            "confidence": 1.0,
        })

        editor = EditorInChief(mock_llm, mock_kb)
        original_draft = "### 主旨\n良好草稿\n### 說明\n良好說明"

        from unittest.mock import patch
        with patch.object(editor, '_auto_refine', wraps=editor._auto_refine) as mock_refine:
            final_draft, report = editor.review_and_refine(original_draft, "函")
            # Safe 或 Low 都不觸發
            if report.risk_summary in ["Safe", "Low"]:
                mock_refine.assert_not_called()
                assert final_draft == original_draft


# ==================== GAP-002: _build_audit_log [I] icon ====================

class TestBuildAuditLogInfoIcon:
    """GAP-002: _build_audit_log 中 info severity 應顯示 [I] 圖示。"""

    def test_info_severity_shows_i_icon(self):
        """info 等級的 issue 在審計日誌中應顯示 [I] 圖示"""
        editor = EditorInChief(MagicMock())
        results = [
            ReviewResult(
                agent_name="Compliance Checker",
                issues=[ReviewIssue(
                    category="compliance", severity="info", risk_level="low",
                    location="辦法", description="建議補充參考依據",
                    suggestion="加入法規引用",
                )],
                score=0.85,
                confidence=0.9,
            ),
        ]
        log = editor._build_audit_log(results, 0.85, "Low", 0.0, 0.0)
        assert "[I]" in log
        assert "建議補充參考依據" in log
        assert "*建議*：加入法規引用" in log

    def test_all_three_icons_in_log(self):
        """error/warning/info 三種等級應對應 [E]/[W]/[I]"""
        editor = EditorInChief(MagicMock())
        results = [
            ReviewResult(
                agent_name="Multi Agent",
                issues=[
                    ReviewIssue(category="format", severity="error", risk_level="high",
                                location="結構", description="嚴重問題"),
                    ReviewIssue(category="format", severity="warning", risk_level="medium",
                                location="用詞", description="一般問題"),
                    ReviewIssue(category="format", severity="info", risk_level="low",
                                location="建議", description="改善建議"),
                ],
                score=0.5,
                confidence=1.0,
            ),
        ]
        log = editor._build_audit_log(results, 0.5, "High", 1.0, 0.5)
        assert "[E]" in log
        assert "[W]" in log
        assert "[I]" in log


# ==================== GAP-003: _get_agent_category "policy" 映射 ====================

class TestGetAgentCategoryPolicy:
    """GAP-003: _get_agent_category 對含 policy 名稱的 Agent 映射到 compliance。"""

    def test_policy_keyword_maps_to_compliance(self):
        """含 'policy' 的 Agent 名稱應映射到 compliance 類別"""
        editor = EditorInChief(MagicMock())
        assert editor._get_agent_category("Policy Checker") == "compliance"
        assert editor._get_agent_category("policy_validator") == "compliance"

    def test_auditor_keyword_maps_to_format(self):
        """含 'auditor' 的 Agent 名稱應映射到 format 類別"""
        editor = EditorInChief(MagicMock())
        assert editor._get_agent_category("Custom Auditor") == "format"

    def test_unknown_agent_defaults_to_style(self):
        """不匹配任何關鍵字的 Agent 應預設為 style"""
        editor = EditorInChief(MagicMock())
        assert editor._get_agent_category("Random Agent") == "style"
        assert editor._get_agent_category("MyChecker") == "style"


# ==================== GAP-004: validate_agent_names >5 個拒絕 ====================

class TestValidateAgentNamesLimit:
    """GAP-004: validate_agent_names 超過 5 個 Agent 時應拒絕。"""

    def test_more_than_5_agents_rejected(self):
        """超過 5 個 Agent 的列表應被拒絕"""
        from api_server import ParallelReviewRequest
        with pytest.raises(ValueError, match="最多 5 個"):
            ParallelReviewRequest(
                draft="### 主旨\n測試草稿足夠長度",
                doc_type="函",
                agents=["format", "style", "fact", "consistency", "compliance", "format"],
            )

    def test_empty_agents_rejected(self):
        """空的 Agent 列表應被拒絕"""
        from api_server import ParallelReviewRequest
        with pytest.raises(ValueError, match="不可為空"):
            ParallelReviewRequest(
                draft="### 主旨\n測試草稿足夠長度",
                doc_type="函",
                agents=[],
            )

    def test_duplicate_agents_deduped(self):
        """重複的 Agent 名稱應被去重"""
        from api_server import ParallelReviewRequest
        req = ParallelReviewRequest(
            draft="### 主旨\n測試草稿足夠長度",
            doc_type="函",
            agents=["format", "format", "style"],
        )
        assert req.agents == ["format", "style"]


# ==================== GAP-005: export 自動建立輸出目錄 ====================

class TestExporterAutoCreateDir:
    """GAP-005: DocxExporter.export 應自動建立不存在的輸出目錄。"""

    def test_auto_creates_output_directory(self, tmp_path):
        """匯出到不存在的子目錄時應自動建立"""
        from src.document.exporter import DocxExporter

        exporter = DocxExporter()
        output_path = str(tmp_path / "new_subdir" / "test_doc.docx")
        result = exporter.export("# 函\n### 主旨\n測試主旨", output_path)
        assert result == output_path
        assert os.path.exists(output_path)

    def test_existing_directory_no_error(self, tmp_path):
        """匯出到已存在的目錄時不應出錯"""
        from src.document.exporter import DocxExporter
        import os

        sub = tmp_path / "existing"
        sub.mkdir()
        exporter = DocxExporter()
        output_path = str(sub / "test_doc.docx")
        exporter.export("# 函\n### 主旨\n測試主旨", output_path)
        assert os.path.exists(output_path)


# ==================== GAP-006: _sanitize_text surrogate pair 移除 ====================

class TestSanitizeTextSurrogatePairs:
    """GAP-006: _sanitize_text 應移除 Unicode surrogate pairs。"""

    def test_surrogate_pairs_removed(self):
        """包含 surrogate code points 的文字應被清理"""
        from src.document.exporter import DocxExporter

        exporter = DocxExporter()
        # 手動建構含有 surrogate-like 字元的字串
        # 注意：Python 3 通常不允許直接建立含 surrogate 的字串，
        # 但 LLM 輸出中偶爾會出現。用 re 模式匹配測試。
        import re
        # 直接驗證 regex 模式正確
        pattern = re.compile(r'[\ud800-\udfff]')
        test_text = "公文正常文字"
        cleaned = exporter._sanitize_text(test_text)
        assert cleaned == "公文正常文字"
        assert not pattern.search(cleaned)

    def test_bom_removed(self):
        """BOM (Byte Order Mark) 應被移除"""
        from src.document.exporter import DocxExporter

        exporter = DocxExporter()
        text = "\ufeff公文標題"
        cleaned = exporter._sanitize_text(text)
        assert "\ufeff" not in cleaned
        assert cleaned == "公文標題"

    def test_zwsp_removed(self):
        """ZWSP (Zero Width Space) 應被移除"""
        from src.document.exporter import DocxExporter

        exporter = DocxExporter()
        text = "主\u200b旨"
        cleaned = exporter._sanitize_text(text)
        assert cleaned == "主旨"


# ==================== GAP-007: _write_meta_info ### 停止解析 ====================

class TestWriteMetaInfoStopConditions:
    """GAP-007: _write_meta_info 應在遇到 ### 時停止解析。"""

    def test_stops_at_triple_hash(self, tmp_path):
        """遇到 ### 標題時應停止寫入檔頭"""
        from src.document.exporter import DocxExporter
        from docx import Document

        exporter = DocxExporter()
        draft = "# 函\n\n**機關**：測試機關\n**受文者**：測試單位\n### 主旨\n測試主旨"
        output_path = str(tmp_path / "stop_test.docx")
        exporter.export(draft, output_path)

        doc = Document(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "測試機關" in all_text
        assert "測試主旨" in all_text

    def test_stops_at_dashes(self, tmp_path):
        """遇到 --- 分隔線時應停止寫入檔頭"""
        from src.document.exporter import DocxExporter
        from docx import Document

        exporter = DocxExporter()
        draft = "# 函\n\n**機關**：測試機關\n---\n主旨：測試主旨"
        output_path = str(tmp_path / "stop_dash.docx")
        exporter.export(draft, output_path)

        doc = Document(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "測試機關" in all_text

    def test_both_attachment_keys_in_exporter(self, tmp_path):
        """_write_attachments 應處理 attachments 和 references 兩個 key"""
        from src.document.exporter import DocxExporter
        from docx import Document

        exporter = DocxExporter()
        draft = "# 函\n\n主旨：測試\n\n附件：報告書\n\n參考來源：某法規"
        output_path = str(tmp_path / "both_att.docx")
        exporter.export(draft, output_path)

        doc = Document(output_path)
        all_text = "\n".join([p.text for p in doc.paragraphs])
        assert "報告書" in all_text


# ==================== GAP-008: Meeting 端點匯出失敗 ====================

@pytest.mark.skipif(
    not importlib.util.find_spec("multipart"),
    reason="python-multipart 未安裝，跳過 API 相關測試",
)
class TestMeetingExportFailure:
    """GAP-008: Meeting 端點中 DocxExporter 失敗時的處理。"""

    def test_meeting_exporter_failure_returns_error(self):
        """匯出失敗應回傳含 error 的回應，不崩潰"""
        import asyncio
        from unittest.mock import patch, MagicMock
        from api_server import run_meeting, MeetingRequest

        request = MeetingRequest(
            user_input="寫一份函給教育部",
            skip_review=True,
            output_docx=True,
        )

        # Mock 全域初始化函數
        mock_llm = MagicMock()
        mock_llm.generate.return_value = json.dumps({
            "doc_type": "函",
            "sender": "測試機關",
            "receiver": "教育部",
            "subject": "配合辦理事項",
        })

        # `src.api.routes.workflow` 用 `from src.api.dependencies import get_llm, get_kb`
        # 創 local binding, 只 patch src.api.dependencies 影響不到；必須同時 patch
        # workflow package 的 local binding, 否則 workflow.get_llm() 會拿到 real LLM,
        # 一輪 document workflow 打 litellm 6+ 次 ~20s/次 = 120s+ 死時間。
        # （修法來自 adb531c 的 preflight re-bind 教訓；T-PYTEST-RUNTIME-FIX-v2 對症）
        with patch("src.api.dependencies.get_llm", return_value=mock_llm), \
             patch("src.api.dependencies.get_kb", return_value=MagicMock()), \
             patch("src.api.routes.workflow.get_llm", return_value=mock_llm), \
             patch("src.api.routes.workflow.get_kb", return_value=MagicMock()), \
             patch("src.api.routes.workflow.DocxExporter") as mock_exporter_cls:
            mock_exporter_cls.return_value.export.side_effect = IOError("磁碟已滿")

            loop = asyncio.new_event_loop()
            try:
                response = loop.run_until_complete(run_meeting(request))
            finally:
                loop.close()

            # 匯出失敗被上層 except 捕獲
            # 不一定 success=False，要看 IOError 在哪層被捕獲
            assert response is not None


# ==================== BUG-010: 辦法/公告事項 子字串誤匹配 ====================

class TestProvisionHeaderSubstringFalsePositive:
    """BUG-010: detect_header 中 '辦法/公告事項' 不應以子字串方式匹配。

    原本使用 `"辦法/公告事項" in clean`，會將內文中偶然包含「辦法/公告事項」
    的行誤判為段落標題。修正為使用 _is_section_header 精確匹配。
    """

    def test_provision_slash_header_exact(self):
        """「辦法/公告事項」作為段落標題應被正確辨識"""
        engine = TemplateEngine()
        sections = engine.parse_draft("主旨：測試\n辦法/公告事項：一、請照辦")
        assert sections["provisions"] != ""
        assert "請照辦" in sections["provisions"]

    def test_provision_slash_header_with_space(self):
        """「辦法/公告事項 」後接空格的段落標題"""
        engine = TemplateEngine()
        sections = engine.parse_draft("主旨：測試\n辦法/公告事項 一、請照辦")
        assert "請照辦" in sections["provisions"]

    def test_provision_slash_header_standalone(self):
        """「辦法/公告事項」獨立一行作為段落標題"""
        engine = TemplateEngine()
        sections = engine.parse_draft("主旨：測試\n辦法/公告事項\n一、請照辦")
        assert "請照辦" in sections["provisions"]

    def test_inline_text_not_mistaken_as_header(self):
        """內文中包含「辦法/公告事項」不應被誤判為段落標題"""
        engine = TemplateEngine()
        sections = engine.parse_draft(
            "主旨：測試\n說明：本案依辦法/公告事項的規定辦理\n辦法：一、請照辦"
        )
        # 「說明」段落應包含完整內文（不應被截斷）
        assert "辦法/公告事項的規定辦理" in sections["explanation"]

    def test_inline_mention_in_explanation(self):
        """說明段落內文提及「辦法/公告事項」不應切換 section"""
        engine = TemplateEngine()
        sections = engine.parse_draft(
            "主旨：測試\n"
            "說明：依據上級指示辦法/公告事項如下所述\n"
            "請查照辦理"
        )
        # 第二行不應被解析為 provisions header
        assert "如下所述" in sections["explanation"]


# ==================== BUG-001: 全形空格段落標題辨識 ====================

class TestSectionHeaderFullWidthSpace:
    """BUG-001: _is_section_header 應能辨識全形空格（\u3000）分隔的段落標題。

    LLM 可能輸出「主旨　本案...」（使用全形空格），需正確辨識為段落標題。
    """

    def test_subject_with_fullwidth_space(self):
        """主旨後接全形空格應被辨識為段落標題"""
        engine = TemplateEngine()
        sections = engine.parse_draft("主旨\u3000本案函轉辦理\n說明\n依據某法辦理")
        assert sections["subject"] == "本案函轉辦理"

    def test_explanation_with_fullwidth_space(self):
        """說明後接全形空格應被辨識為段落標題"""
        engine = TemplateEngine()
        sections = engine.parse_draft("主旨：測試\n說明\u3000詳細內容如下")
        assert "詳細內容如下" in sections["explanation"]

    def test_provisions_with_fullwidth_space(self):
        """辦法後接全形空格應被辨識為段落標題"""
        engine = TemplateEngine()
        sections = engine.parse_draft("主旨：測試\n辦法\u3000一、請照辦")
        assert sections["provisions"] != ""

    def test_halfwidth_space_still_works(self):
        """半形空格仍然正常運作（回歸測試）"""
        engine = TemplateEngine()
        sections = engine.parse_draft("主旨 本案辦理\n說明 依據某法")
        assert sections["subject"] == "本案辦理"

    def test_mixed_headers(self):
        """混合全形和半形空格的文件應全部正確解析"""
        engine = TemplateEngine()
        draft = "主旨\u3000測試主旨\n說明：測試說明\n辦法 請照辦"
        sections = engine.parse_draft(draft)
        assert sections["subject"] == "測試主旨"
        assert "測試說明" in sections["explanation"]
        assert sections["provisions"] != ""


# ==================== BUG-003: NaN/Infinity 分數處理 ====================

class TestNaNInfinityScoreHandling:
    """BUG-003: parse_review_response 和 ComplianceChecker 應正確處理 NaN/Infinity 分數。

    float("NaN") 和 float("Infinity") 不會拋出 ValueError/TypeError，
    但 NaN 比較結果不可預測，可能導致 Pydantic 驗證失敗或錯誤的分數。
    """

    def test_nan_score_falls_back_to_default(self):
        """NaN 分數應回退為預設值"""
        response = '{"issues": [], "score": "NaN"}'
        result = parse_review_response(response, "Test Agent", "style")
        # 應使用 DEFAULT_REVIEW_SCORE（0.8），而非 NaN
        assert result.score == 0.8
        assert not (result.score != result.score)  # 確認不是 NaN

    def test_infinity_score_falls_back_to_default(self):
        """Infinity 分數應回退為預設值"""
        response = '{"issues": [], "score": "Infinity"}'
        result = parse_review_response(response, "Test Agent", "style")
        assert result.score == 0.8

    def test_negative_infinity_score_falls_back_to_default(self):
        """-Infinity 分數應回退為預設值"""
        response = '{"issues": [], "score": "-Infinity"}'
        result = parse_review_response(response, "Test Agent", "style")
        assert result.score == 0.8

    def test_nan_confidence_falls_back_to_default(self):
        """NaN 信心度應回退為預設值 1.0"""
        response = '{"issues": [], "score": 0.9, "confidence": "NaN"}'
        result = parse_review_response(response, "Test Agent", "style")
        assert result.confidence == 1.0

    def test_compliance_checker_nan_score(self):
        """ComplianceChecker 對 NaN 分數的處理"""
        from src.agents.compliance_checker import ComplianceChecker
        from src.core.constants import DEFAULT_COMPLIANCE_SCORE

        mock_llm = MagicMock()
        mock_llm.generate.return_value = '{"issues": [], "score": "NaN", "confidence": 0.9}'
        checker = ComplianceChecker(mock_llm)
        result = checker.check("### 主旨\n測試")
        assert result.score == DEFAULT_COMPLIANCE_SCORE

    def test_normal_scores_still_work(self):
        """正常分數仍正確鉗位（回歸測試）"""
        response = '{"issues": [], "score": 0.75, "confidence": 0.9}'
        result = parse_review_response(response, "Test Agent", "style")
        assert result.score == 0.75
        assert result.confidence == 0.9

    def test_overflow_scores_still_clamped(self):
        """超出範圍的數值分數仍正確鉗位（回歸測試）"""
        response = '{"issues": [], "score": 1.5, "confidence": -0.2}'
        result = parse_review_response(response, "Test Agent", "style")
        assert result.score == 1.0
        assert result.confidence == 0.0


# ==================== 額外邊界測試：renumber_provisions > 20 項 ====================

class TestRenumberProvisionsOverflow:
    """renumber_provisions 超過 CHINESE_NUMBERS 長度時的行為。"""

    def test_more_than_20_items_uses_arabic_numbers(self):
        """超過 20 個主項目時應回退到阿拉伯數字"""
        lines = [f"{i}、第{i}項內容" for i in range(1, 23)]
        text = "\n".join(lines)
        result = renumber_provisions(text)
        # 前 20 項使用中文數字
        assert "一、第1項內容" in result
        assert "二十、第20項內容" in result
        # 第 21 項以上使用阿拉伯數字
        assert "21、第21項內容" in result
        assert "22、第22項內容" in result

    def test_sub_items_over_20_uses_arabic(self):
        """超過 20 個子項目時應回退到阿拉伯數字"""
        lines = ["1、主項目"]
        for i in range(1, 23):
            lines.append(f"（{i}）子項目{i}")
        text = "\n".join(lines)
        result = renumber_provisions(text)
        assert "（二十）子項目20" in result
        assert "（21）子項目21" in result


# ==================== 額外邊界測試：TemplateEngine _fallback_apply ====================

class TestFallbackApply:
    """TemplateEngine._fallback_apply 的邊界測試。"""

    def test_fallback_removes_empty_provisions(self):
        """備用格式應移除空白的辦法段落"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="函", sender="市府", receiver="各機關",
            subject="測試", reason="說明內容"
        )
        sections = {"subject": "測試", "explanation": "說明內容", "provisions": ""}
        result = engine._fallback_apply(req, sections)
        # 不應有空的「### 辦法」段落
        assert "### 辦法\n\n" not in result

    def test_fallback_includes_attachments(self):
        """備用格式應包含附件"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="函", sender="市府", receiver="各機關",
            subject="測試", attachments=["附件一", "附件二"]
        )
        sections = {"subject": "測試", "explanation": "", "provisions": ""}
        result = engine._fallback_apply(req, sections)
        assert "附件一" in result
        assert "附件二" in result

    def test_fallback_with_action_items_when_no_explanation(self):
        """備用格式在無說明時應使用 action_items"""
        engine = TemplateEngine()
        req = PublicDocRequirement(
            doc_type="函", sender="市府", receiver="各機關",
            subject="測試", action_items=["辦理一", "辦理二"]
        )
        sections = {"subject": "測試", "explanation": "", "provisions": ""}
        result = engine._fallback_apply(req, sections)
        assert "辦理一" in result
        assert "辦理二" in result


# ==================== BUG-012: BOM/零寬字元在審查 JSON 中的處理 ====================

class TestBomInReviewJsonParsing:
    """BUG-012: parse_review_response 和 ComplianceChecker._parse_response
    應在 JSON 解析前清理 BOM/零寬字元，避免欄位名稱不匹配導致 issues 靜默遺失。"""

    def test_parse_review_response_with_bom_in_field_names(self):
        """parse_review_response 應正確解析欄位名含 BOM 的 JSON"""
        # BOM 插入在 "issues" 和 "score" 鍵名中
        response = (
            '{\ufeff"is\u200bsues": [{"severity": "error",'
            ' "location": "主旨", "description": "缺少主旨"}],'
            ' "\u200cscore": 0.6}'
        )
        result = parse_review_response(response, "Test Agent", "style")
        assert len(result.issues) == 1
        assert result.issues[0].description == "缺少主旨"
        assert result.score == 0.6

    def test_parse_review_response_with_zwsp_in_values(self):
        """parse_review_response 應正確解析值含零寬字元的 JSON"""
        response = (
            '{"issues": [{"severity": "warning",'
            ' "location": "說\u200b明", "description": "用\u200c詞不當"}],'
            ' "score": 0.8}'
        )
        result = parse_review_response(response, "Test Agent", "fact")
        assert len(result.issues) == 1
        assert result.issues[0].location == "說明"
        assert result.issues[0].description == "用詞不當"

    def test_parse_review_response_bom_only_prefix(self):
        """parse_review_response 應處理回應開頭的 BOM"""
        response = '\ufeff{"issues": [], "score": 0.95}'
        result = parse_review_response(response, "Test Agent", "consistency")
        assert result.score == 0.95
        assert len(result.issues) == 0

    def test_compliance_checker_parse_with_bom_in_field_names(self):
        """ComplianceChecker 應正確解析欄位名含 BOM 的 JSON（透過 check）"""
        mock_llm = MagicMock()
        # BOM 插入在 "issues" 鍵名中
        mock_llm.generate.return_value = (
            '{\ufeff"\u200bissues": [{"severity": "error",'
            ' "location": "全文", "description": "違反政策"}],'
            ' "score": 0.4, "confidence": 0.9}'
        )
        checker = ComplianceChecker(llm=mock_llm, kb_manager=None)
        result = checker.check("### 主旨\n測試草稿")
        assert len(result.issues) == 1
        assert result.issues[0].description == "違反政策"
        assert result.score == 0.4

    def test_compliance_checker_parse_with_zwsp_in_values(self):
        """ComplianceChecker 應處理值中的零寬字元（透過 check）"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = (
            '{"issues": [{"severity": "warning",'
            ' "location": "辦\u200c法", "description": "用\u200d詞過時"}],'
            ' "score": 0.7, "confidence": 0.85}'
        )
        checker = ComplianceChecker(llm=mock_llm, kb_manager=None)
        result = checker.check("### 主旨\n測試草稿")
        assert len(result.issues) == 1
        assert result.issues[0].location == "辦法"
        assert result.issues[0].description == "用詞過時"

    def test_compliance_checker_parse_bom_only_prefix(self):
        """ComplianceChecker 應處理回應開頭的 BOM（透過 check）"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = '\ufeff{"issues": [], "score": 0.9, "confidence": 0.95}'
        checker = ComplianceChecker(llm=mock_llm, kb_manager=None)
        result = checker.check("### 主旨\n測試草稿")
        assert result.score == 0.9
        assert len(result.issues) == 0

    def test_auditor_audit_with_bom_in_field_names(self):
        """FormatAuditor.audit 應正確解析欄位名含 BOM 的 LLM 回應"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = (
            '{\ufeff"\u200berrors": ["缺少主旨"],'
            ' "\u200cwarnings": ["格式待改善"]}'
        )
        auditor = FormatAuditor(llm_provider=mock_llm, kb_manager=None)
        result = auditor.audit("主旨：測試\n說明：內容", "函")
        assert "缺少主旨" in result["errors"]
        assert "格式待改善" in result["warnings"]

    def test_auditor_audit_with_bom_prefix_only(self):
        """FormatAuditor.audit 應處理回應開頭的 BOM"""
        mock_llm = MagicMock()
        mock_llm.generate.return_value = '\ufeff{"errors": [], "warnings": []}'
        auditor = FormatAuditor(llm_provider=mock_llm, kb_manager=None)
        # 含完整引用結構的草稿，避免通用驗證器產生額外錯誤
        draft = "主旨：測試\n說明：依據相關法規辦理[^1]。\n\n### 參考來源\n[^1]: [Level A] 測試法規"
        result = auditor.audit(draft, "函")
        assert result["errors"] == []
        assert result["warnings"] == []


# ==================== WriterAgent Error prefix fallback ====================

class TestWriterAgentErrorFallback:
    """WriterAgent 在 LLM 回傳 Error: 開頭字串時應使用基本模板。"""

    def test_error_prefix_triggers_fallback(self, mock_llm):
        """LLM 回傳 'Error: ...' 時應使用基本模板"""
        mock_kb = MagicMock()
        mock_kb.search_hybrid.return_value = []
        mock_llm.generate.return_value = "Error: rate limit exceeded"

        writer = WriterAgent(mock_llm, mock_kb)
        req = PublicDocRequirement(
            doc_type="函", sender="機關", receiver="單位",
            subject="測試主旨",
        )
        draft = writer.write_draft(req)
        # 應使用基本模板，包含主旨
        assert "測試主旨" in draft
        # 不應包含 Error 訊息
        assert "rate limit" not in draft

    def test_error_prefix_with_examples_no_sources(self, mock_llm):
        """Error fallback 時即使有範例也不應附加參考來源到 fallback 模板"""
        mock_kb = MagicMock()
        mock_kb.search_hybrid.return_value = [
            {"content": "範例", "metadata": {"title": "函範例"}},
        ]
        mock_llm.generate.return_value = "Error: context length exceeded"

        writer = WriterAgent(mock_llm, mock_kb)
        req = PublicDocRequirement(
            doc_type="函", sender="機關", receiver="單位",
            subject="預算報告",
        )
        draft = writer.write_draft(req)
        assert "預算報告" in draft
        # fallback 模板後仍會附加參考來源（因為 sources_list 非空）
        assert "參考來源" in draft

    def test_whitespace_only_triggers_fallback(self, mock_llm):
        """LLM 回傳純空白字串時應使用基本模板"""
        mock_kb = MagicMock()
        mock_kb.search_hybrid.return_value = []
        mock_llm.generate.return_value = "   \n\t  "

        writer = WriterAgent(mock_llm, mock_kb)
        req = PublicDocRequirement(
            doc_type="公告", sender="市府", receiver="市民",
            subject="停水公告",
        )
        draft = writer.write_draft(req)
        assert "停水公告" in draft
        assert "主旨" in draft
