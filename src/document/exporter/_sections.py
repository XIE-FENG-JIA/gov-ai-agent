from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.agents.template import clean_markdown_artifacts
from src.core.constants import FIRST_LINE_INDENT, FONT_LOG, FONT_SIZE_LOG

_BODY_START_KEYWORDS = ("主旨", "通話時間", "發話人", "指示事項")


def _body_order_for(doc_type: str) -> list[tuple[str, str]]:
    if doc_type == "公告":
        return [("主旨", "subject"), ("依據", "basis"), ("公告事項", "provisions")]
    if doc_type == "簽":
        return [("主旨", "subject"), ("說明", "explanation"), ("擬辦", "provisions")]
    if doc_type == "開會通知單":
        return [
            ("主旨", "subject"),
            ("說明", "explanation"),
            ("開會時間", "meeting_time"),
            ("開會地點", "meeting_location"),
            ("議程", "agenda"),
            ("注意事項", "provisions"),
        ]
    if doc_type == "會勘通知單":
        return [
            ("主旨", "subject"),
            ("說明", "explanation"),
            ("會勘時間", "inspection_time"),
            ("會勘地點", "inspection_location"),
            ("會勘事項", "inspection_items"),
            ("應攜文件", "required_documents"),
            ("應出席單位", "attendees"),
            ("注意事項", "provisions"),
        ]
    if doc_type == "公務電話紀錄":
        return [
            ("通話時間", "call_time"),
            ("發話人", "caller"),
            ("受話人", "callee"),
            ("主旨", "subject"),
            ("通話摘要", "call_summary"),
            ("說明", "explanation"),
            ("追蹤事項", "follow_up_items"),
            ("紀錄人", "recorder"),
            ("核閱", "reviewer"),
        ]
    if doc_type == "手令":
        return [
            ("主旨", "subject"),
            ("指示事項", "directive_content"),
            ("說明", "explanation"),
            ("完成期限", "deadline"),
            ("副知", "cc_list"),
        ]
    if doc_type == "開會紀錄":
        return [
            ("主席（主持人）", "chairperson"),
            ("出席人員", "attendees"),
            ("列席人員", "observers"),
            ("請假人員", "absentees"),
            ("紀錄", "recorder"),
            ("報告事項", "report_items"),
            ("討論事項", "discussion_items"),
            ("決議", "resolutions"),
            ("臨時動議", "motions"),
            ("主席結論", "chairman_conclusion"),
            ("散會時間", "adjournment_time"),
        ]
    if doc_type == "箋函":
        return [("主旨", "subject"), ("說明", "explanation"), ("正本", "copies_to"), ("副本", "cc_copies")]
    return [("主旨", "subject"), ("說明", "explanation"), ("辦法", "provisions")]


def write_title(exporter, doc, doc_type: str) -> None:
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exporter._set_paragraph_spacing(p_title)
    run = p_title.add_run(doc_type)
    exporter._set_font(run, exporter._font_title, exporter._size_title, bold=True)


def write_meta_info(exporter, doc, draft_text: str) -> None:
    lines = draft_text.strip().split("\n")
    for line in lines[1:]:
        clean = exporter._clean_line(line)
        if not clean:
            continue
        if clean.startswith("---") or clean.startswith("###"):
            break
        if any(clean.startswith(keyword) for keyword in _BODY_START_KEYWORDS):
            break
        if clean in exporter.KNOWN_DOC_TYPES:
            continue
        if clean:
            paragraph = doc.add_paragraph()
            exporter._set_paragraph_spacing(paragraph)
            run = paragraph.add_run(clean)
            exporter._set_font(run, exporter._font_body, exporter._size_meta)
    doc.add_paragraph()


def write_body(exporter, doc, sections: dict[str, str], doc_type: str = "函") -> None:
    for label, key in _body_order_for(doc_type):
        content = sections.get(key)
        if not content:
            continue

        p_label = doc.add_paragraph()
        exporter._set_paragraph_spacing(p_label)
        run_label = p_label.add_run(f"{label}：")
        exporter._set_font(run_label, exporter._font_title, exporter._size_section, bold=True)

        lines = exporter._auto_number(content.split("\n")) if exporter.strict_format else content.split("\n")
        for line in lines:
            clean_content = exporter._clean_line(line)
            if clean_content:
                paragraph = doc.add_paragraph()
                exporter._set_paragraph_spacing(paragraph)
                paragraph.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT)
                run_content = paragraph.add_run(clean_content)
                exporter._set_font(run_content, exporter._font_body, exporter._size_body)


def write_attachments(exporter, doc, sections: dict[str, str]) -> None:
    for key, label in [("attachments", "附件："), ("references", "參考來源：")]:
        content = sections.get(key)
        if not content:
            continue

        doc.add_paragraph()
        p_att = doc.add_paragraph()
        exporter._set_paragraph_spacing(p_att)
        run_att = p_att.add_run(label)
        exporter._set_font(run_att, exporter._font_title, exporter._size_body, bold=True)

        for line in content.split("\n"):
            clean_content = exporter._clean_line(line)
            if clean_content:
                paragraph = doc.add_paragraph(clean_content)
                exporter._set_paragraph_spacing(paragraph)
                for run in paragraph.runs:
                    exporter._set_font(run, exporter._font_body, exporter._size_meta)


def write_qa_report(exporter, doc, qa_report: str) -> None:
    doc.add_page_break()
    p_qa_title = doc.add_paragraph()
    exporter._set_paragraph_spacing(p_qa_title)
    run_qa = p_qa_title.add_run("附件：AI 品質保證報告 (QA Report)")
    exporter._set_font(run_qa, exporter._font_title, exporter._size_section, bold=True)

    clean_qa = exporter._sanitize_text(qa_report)
    clean_qa = clean_markdown_artifacts(clean_qa)
    paragraph = doc.add_paragraph(clean_qa)
    for run in paragraph.runs:
        run.font.name = FONT_LOG
        run.font.size = Pt(FONT_SIZE_LOG)
