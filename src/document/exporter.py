import logging
import os
import re
import tempfile
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from src.agents.template import TemplateEngine, clean_markdown_artifacts
from src.document.citation_metadata import build_citation_export_metadata
from src.core.constants import (
    FONT_SIZE_TITLE,
    FONT_SIZE_SECTION_LABEL,
    FONT_SIZE_BODY,
    FONT_SIZE_META,
    FONT_SIZE_LOG,
    STRICT_FONT_SIZE_TITLE,
    STRICT_FONT_SIZE_SECTION_LABEL,
    STRICT_FONT_SIZE_BODY,
    STRICT_FONT_SIZE_META,
    STRICT_PAGE_MARGIN,
    STRICT_LINE_SPACING,
    STRICT_SPACE_BEFORE_LINES,
    STRICT_SPACE_AFTER_LINES,
    PAGE_MARGIN_TOP,
    PAGE_MARGIN_BOTTOM,
    PAGE_MARGIN_LEFT,
    PAGE_MARGIN_RIGHT,
    FIRST_LINE_INDENT,
    FONT_LOG,
    CHINESE_NUMBERS,
    get_platform_fonts,
)

logger = logging.getLogger(__name__)

# 所有已知的公文類型（用於識別標題行）
KNOWN_DOC_TYPES = frozenset([
    "函", "公告", "簽", "書函", "令", "開會通知單", "開會紀錄",
    "呈", "咨", "會勘通知單", "公務電話紀錄", "手令", "箋函",
])

CUSTOM_PROPERTY_FMTID = "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}"
CUSTOM_PROPERTY_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties"
CUSTOM_PROPERTY_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.custom-properties+xml"
CONTENT_TYPES_XML_PATH = "[Content_Types].xml"
PACKAGE_RELS_XML_PATH = "_rels/.rels"
CUSTOM_PROPERTIES_NS = "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
DOC_PROPS_VT_NS = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PACKAGE_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CUSTOM_PROPERTIES_XML_PATH = "docProps/custom.xml"

ET.register_namespace("cp", CUSTOM_PROPERTIES_NS)
ET.register_namespace("vt", DOC_PROPS_VT_NS)
ET.register_namespace("", CONTENT_TYPES_NS)
ET.register_namespace("", PACKAGE_RELS_NS)


class DocxExporter:
    """
    將 Markdown 草稿匯出為符合政府標準格式的 Microsoft Word 文件。
    """

    def __init__(self, strict_format: bool = True) -> None:
        self.template_engine = TemplateEngine()
        self.strict_format = strict_format
        # 跨平台字體：依據作業系統選擇適當的中文字體
        title_font, body_font = get_platform_fonts()
        self._font_title: str | None = title_font
        self._font_body: str | None = body_font
        # 依嚴格模式選擇字型大小
        if strict_format:
            self._size_title = STRICT_FONT_SIZE_TITLE
            self._size_section = STRICT_FONT_SIZE_SECTION_LABEL
            self._size_body = STRICT_FONT_SIZE_BODY
            self._size_meta = STRICT_FONT_SIZE_META
        else:
            self._size_title = FONT_SIZE_TITLE
            self._size_section = FONT_SIZE_SECTION_LABEL
            self._size_body = FONT_SIZE_BODY
            self._size_meta = FONT_SIZE_META

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """
        清理文字中可能導致 Word 文件損壞的特殊字元。

        移除 XML 不允許的控制字元（保留換行和 Tab），
        並替換可能導致問題的 Unicode 字元。
        """
        if not text:
            return ""
        # 移除 XML 1.0 不允許的控制字元（0x00-0x08, 0x0B-0x0C, 0x0E-0x1F）
        # 保留 0x09 (Tab), 0x0A (LF), 0x0D (CR)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        # 移除 Unicode 替代碼位（surrogate pairs，可能出現在錯誤的 LLM 輸出中）
        text = re.sub(r'[\ud800-\udfff]', '', text)
        # 替換 Unicode 特殊空格為普通空格
        for ch in [
            '\u00a0',   # NBSP (No-Break Space)
            '\u2000',   # EN Quad
            '\u2001',   # EM Quad
            '\u2002',   # EN Space
            '\u2003',   # EM Space
            '\u2004',   # Three-Per-Em Space
            '\u2005',   # Four-Per-Em Space
            '\u2006',   # Six-Per-Em Space
            '\u2007',   # Figure Space
            '\u2008',   # Punctuation Space
            '\u2009',   # Thin Space
            '\u200a',   # Hair Space
            '\u202f',   # Narrow No-Break Space
            '\u205f',   # Medium Mathematical Space
            '\u3000',   # Ideographic Space（全形空格）
        ]:
            text = text.replace(ch, ' ')
        # 移除零寬度字元和其他不可見 Unicode 字元
        for ch in [
            '\ufeff',   # BOM (Byte Order Mark)
            '\u200b',   # ZWSP (Zero Width Space)
            '\u200c',   # ZWNJ (Zero Width Non-Joiner)
            '\u200d',   # ZWJ (Zero Width Joiner)
            '\u200e',   # LRM (Left-to-Right Mark)
            '\u200f',   # RLM (Right-to-Left Mark)
            '\u00ad',   # Soft Hyphen
            '\u2060',   # Word Joiner
        ]:
            text = text.replace(ch, '')
        return text

    def _set_paragraph_spacing(self, paragraph) -> None:
        """設定段落的行距與段前/段後間距（嚴格模式）。"""
        if not self.strict_format:
            return
        pf = paragraph.paragraph_format
        pf.line_spacing = STRICT_LINE_SPACING
        # 段前 0.5 行：以 Pt 為單位，0.5 行 ≈ 0.5 × body_font_size
        pf.space_before = Pt(self._size_body * STRICT_SPACE_BEFORE_LINES)
        pf.space_after = Pt(STRICT_SPACE_AFTER_LINES)

    def _set_font(self, run, font_name: str | None = None, size: int = FONT_SIZE_META, bold: bool = False) -> None:
        """設定文字段落的字體、大小和粗體。

        若 font_name 為 None，則不設定字體名稱，交由 Word 使用預設字體。
        """
        run.font.size = Pt(size)
        run.font.bold = bold
        if font_name is not None:
            run.font.name = font_name
            # 確保 rPr 元素存在後再設定東亞字體
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn("w:eastAsia"), font_name)

    def _clean_line(self, line: str) -> str:
        """清理單行文字中的 Markdown 標記和特殊字元。"""
        line = self._sanitize_text(line)
        line = clean_markdown_artifacts(line)
        # 額外清理：移除行首的 Markdown 標記殘留
        line = re.sub(r"^[#\*\-]+\s*", "", line)
        return line.strip()

    def _extract_doc_type(self, draft_text: str) -> str:
        """從草稿前 10 行中提取公文類型。

        先精確匹配整行，再嘗試在行中搜尋已知類型關鍵字，
        以處理「# 臺北市政府公告」等格式。
        """
        lines = draft_text.strip().split("\n")
        header_scan_lines = 10

        # 第一遍：精確匹配（清理後整行等於類型名稱）
        for line in lines[:header_scan_lines]:
            clean = self._clean_line(line)
            if clean in KNOWN_DOC_TYPES:
                return clean

        # 第二遍：模糊匹配（行中包含類型關鍵字，優先長名稱）
        # 單字類型（令、呈、咨、函、簽）須出現在行尾才匹配，
        # 避免「茲令各局處」「呈報」等一般文句造成誤判，
        # 但允許「市府內部簽」「行政院令」等標題行正確匹配
        for line in lines[:header_scan_lines]:
            clean = self._clean_line(line)
            for doc_type in sorted(KNOWN_DOC_TYPES, key=len, reverse=True):
                if len(doc_type) == 1:
                    if clean.endswith(doc_type):
                        return doc_type
                    continue
                if doc_type in clean:
                    return doc_type

        return "函"  # 預設公文類型

    @staticmethod
    def _build_custom_properties_xml(properties: dict[str, object]) -> bytes:
        root = ET.Element(f"{{{CUSTOM_PROPERTIES_NS}}}Properties")
        for pid, (name, value) in enumerate(properties.items(), start=2):
            prop = ET.SubElement(
                root,
                f"{{{CUSTOM_PROPERTIES_NS}}}property",
                {
                    "fmtid": CUSTOM_PROPERTY_FMTID,
                    "pid": str(pid),
                    "name": name,
                },
            )
            if isinstance(value, bool):
                child = ET.SubElement(prop, f"{{{DOC_PROPS_VT_NS}}}bool")
                child.text = "true" if value else "false"
            elif isinstance(value, int):
                child = ET.SubElement(prop, f"{{{DOC_PROPS_VT_NS}}}i4")
                child.text = str(value)
            else:
                child = ET.SubElement(prop, f"{{{DOC_PROPS_VT_NS}}}lpwstr")
                child.text = str(value)
        return ET.tostring(root, encoding="UTF-8", xml_declaration=True)

    @staticmethod
    def _ensure_content_types_custom_override(content_types_xml: bytes) -> bytes:
        root = ET.fromstring(content_types_xml)
        override_xpath = f"{{{CONTENT_TYPES_NS}}}Override"
        exists = any(
            node.get("PartName") == f"/{CUSTOM_PROPERTIES_XML_PATH}"
            for node in root.findall(override_xpath)
        )
        if not exists:
            ET.SubElement(
                root,
                override_xpath,
                {
                    "PartName": f"/{CUSTOM_PROPERTIES_XML_PATH}",
                    "ContentType": CUSTOM_PROPERTY_CONTENT_TYPE,
                },
            )
        return ET.tostring(root, encoding="UTF-8", xml_declaration=True)

    @staticmethod
    def _ensure_package_relationship_custom_props(rels_xml: bytes) -> bytes:
        root = ET.fromstring(rels_xml)
        rel_xpath = f"{{{PACKAGE_RELS_NS}}}Relationship"
        exists = any(
            node.get("Type") == CUSTOM_PROPERTY_REL_TYPE and node.get("Target") == "docProps/custom.xml"
            for node in root.findall(rel_xpath)
        )
        if not exists:
            existing_ids = {node.get("Id", "") for node in root.findall(rel_xpath)}
            next_id = 1
            while f"rId{next_id}" in existing_ids:
                next_id += 1
            ET.SubElement(
                root,
                rel_xpath,
                {
                    "Id": f"rId{next_id}",
                    "Type": CUSTOM_PROPERTY_REL_TYPE,
                    "Target": "docProps/custom.xml",
                },
            )
        return ET.tostring(root, encoding="UTF-8", xml_declaration=True)

    @classmethod
    def _write_custom_properties(cls, output_path: str, properties: dict[str, object]) -> None:
        with zipfile.ZipFile(output_path, "r") as zin:
            entries = {name: zin.read(name) for name in zin.namelist()}

        entries[CUSTOM_PROPERTIES_XML_PATH] = cls._build_custom_properties_xml(properties)
        entries[CONTENT_TYPES_XML_PATH] = cls._ensure_content_types_custom_override(entries[CONTENT_TYPES_XML_PATH])
        entries[PACKAGE_RELS_XML_PATH] = cls._ensure_package_relationship_custom_props(entries[PACKAGE_RELS_XML_PATH])

        output_dir = os.path.dirname(output_path) or "."
        fd, temp_path = tempfile.mkstemp(suffix=".docx", dir=output_dir)
        os.close(fd)
        try:
            with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for name, content in entries.items():
                    zout.writestr(name, content)
            os.replace(temp_path, output_path)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

    def export(
        self,
        draft_text: str,
        output_path: str,
        qa_report: str | None = None,
        citation_metadata: dict | None = None,
    ) -> str:
        """將 Markdown 草稿轉換為 docx 檔案。"""
        # 防護空值輸入
        if not draft_text or not draft_text.strip():
            logger.warning("匯出收到空的草稿，產生空白文件")
            draft_text = "（無內容）"

        # 先清理整體文字和特殊字元
        draft_text = self._sanitize_text(draft_text)
        draft_text = clean_markdown_artifacts(draft_text)

        doc = Document()
        self._setup_page(doc)

        # 解析內容結構（失敗時回退為空結構，產生最小有效 DOCX）
        try:
            sections = self.template_engine.parse_draft(draft_text)
        except Exception as exc:
            logger.warning("草稿解析失敗，產生最小有效文件: %s", exc)
            sections = {}

        # 偵測公文類型（供本文段落標籤使用）
        doc_type = self._extract_doc_type(draft_text)

        # 1. 公文類型標題
        self._write_title(doc, doc_type)

        # 2. 檔頭資訊（機關、受文者等）
        self._write_meta_info(doc, draft_text)

        # 3. 本文（主旨、說明、辦法 — 依文件類型調整標籤）
        self._write_body(doc, sections, doc_type)

        # 4. 附件與參考來源
        self._write_attachments(doc, sections)

        # 5. QA 報告（選用附件）
        if qa_report:
            self._write_qa_report(doc, qa_report)

        # 確保輸出目錄存在
        output_dir = os.path.dirname(output_path)
        if output_dir:
            try:
                os.makedirs(output_dir, exist_ok=True)
            except OSError as exc:
                logger.error("建立輸出目錄失敗（%s）: %s", output_dir, exc)
                raise

        try:
            doc.save(output_path)
        except OSError as exc:
            logger.error("DOCX 儲存失敗（%s）: %s", output_path, exc)
            raise

        export_metadata = build_citation_export_metadata(draft_text, citation_metadata)
        if export_metadata:
            self._write_custom_properties(output_path, export_metadata)
        return output_path

    def _setup_page(self, doc: Document) -> None:
        """設定 A4 頁面邊距。"""
        section = doc.sections[0]
        if self.strict_format:
            section.top_margin = Cm(STRICT_PAGE_MARGIN)
            section.bottom_margin = Cm(STRICT_PAGE_MARGIN)
            section.left_margin = Cm(STRICT_PAGE_MARGIN)
            section.right_margin = Cm(STRICT_PAGE_MARGIN)
        else:
            section.top_margin = Cm(PAGE_MARGIN_TOP)
            section.bottom_margin = Cm(PAGE_MARGIN_BOTTOM)
            section.left_margin = Cm(PAGE_MARGIN_LEFT)
            section.right_margin = Cm(PAGE_MARGIN_RIGHT)

    def _write_title(self, doc: Document, doc_type: str) -> None:
        """寫入公文類型標題。"""
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._set_paragraph_spacing(p_title)
        run = p_title.add_run(doc_type)
        self._set_font(run, self._font_title, self._size_title, bold=True)

    def _write_meta_info(self, doc: Document, draft_text: str) -> None:
        """寫入檔頭資訊區塊。"""
        lines = draft_text.strip().split("\n")
        for line in lines[1:]:
            clean = self._clean_line(line)
            if not clean:
                continue
            # 遇到本文段落標籤則結束檔頭區塊
            _BODY_START_KEYWORDS = ("主旨", "通話時間", "發話人", "指示事項")
            if clean.startswith("---") or clean.startswith("###"):
                break
            if any(clean.startswith(kw) for kw in _BODY_START_KEYWORDS):
                break
            # 跳過公文類型重複
            if clean in KNOWN_DOC_TYPES:
                continue
            if clean:
                p = doc.add_paragraph()
                self._set_paragraph_spacing(p)
                run = p.add_run(clean)
                self._set_font(run, self._font_body, self._size_meta)

        doc.add_paragraph()  # 間距

    def _write_body(self, doc: Document, sections: dict[str, str], doc_type: str = "函") -> None:
        """寫入主旨、說明、辦法等本文段落（依文件類型調整標籤）。"""
        if doc_type == "公告":
            # 公告使用「依據」和「公告事項」，不使用「說明」和「辦法」
            body_order = [("主旨", "subject"), ("依據", "basis"), ("公告事項", "provisions")]
        elif doc_type == "簽":
            # 簽使用「擬辦」而非「辦法」
            body_order = [("主旨", "subject"), ("說明", "explanation"), ("擬辦", "provisions")]
        elif doc_type == "開會通知單":
            body_order = [
                ("主旨", "subject"),
                ("說明", "explanation"),
                ("開會時間", "meeting_time"),
                ("開會地點", "meeting_location"),
                ("議程", "agenda"),
                ("注意事項", "provisions"),
            ]
        elif doc_type == "會勘通知單":
            body_order = [
                ("主旨", "subject"),
                ("說明", "explanation"),
                ("會勘時間", "inspection_time"),
                ("會勘地點", "inspection_location"),
                ("會勘事項", "inspection_items"),
                ("應攜文件", "required_documents"),
                ("應出席單位", "attendees"),
                ("注意事項", "provisions"),
            ]
        elif doc_type == "公務電話紀錄":
            body_order = [
                ("通話時間", "call_time"),
                ("發話人", "caller"),
                ("受話人", "callee"),
                ("主旨", "subject"),
                ("通話摘要", "call_summary"),
                ("說明", "explanation"),
                ("追蹤事項", "follow_up_items"),
                ("紀錄人", "recorder"),
                ("核閱", "reviewer"),
            ]
        elif doc_type == "手令":
            body_order = [
                ("主旨", "subject"),
                ("指示事項", "directive_content"),
                ("說明", "explanation"),
                ("完成期限", "deadline"),
                ("副知", "cc_list"),
            ]
        elif doc_type == "開會紀錄":
            body_order = [
                ("主席（主持人）", "chairperson"),
                ("出席人員", "attendees"),
                ("列席人員", "observers"),
                ("請假人員", "absentees"),
                ("紀錄", "recorder"),
                ("報告事項", "report_items"),
                ("討論事項", "discussion_items"),
                ("決議", "resolutions"),
                ("臨時動議", "motions"),
                ("主席結論", "chairman_conclusion"),
                ("散會時間", "adjournment_time"),
            ]
        elif doc_type == "箋函":
            body_order = [
                ("主旨", "subject"),
                ("說明", "explanation"),
                ("正本", "copies_to"),
                ("副本", "cc_copies"),
            ]
        else:
            body_order = [("主旨", "subject"), ("說明", "explanation"), ("辦法", "provisions")]

        for label, key in body_order:
            content = sections.get(key)
            if not content:
                continue

            # 段落標籤
            p_label = doc.add_paragraph()
            self._set_paragraph_spacing(p_label)
            run_label = p_label.add_run(f"{label}：")
            self._set_font(run_label, self._font_title, self._size_section, bold=True)

            # 段落內容（嚴格模式啟用自動編號）
            lines = content.split("\n")
            if self.strict_format:
                lines = self._auto_number(lines)
            for line in lines:
                clean_content = self._clean_line(line)
                if clean_content:
                    p_content = doc.add_paragraph()
                    self._set_paragraph_spacing(p_content)
                    p_content.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT)
                    run_content = p_content.add_run(clean_content)
                    self._set_font(run_content, self._font_body, self._size_body)

    # ── 自動編號 ──────────────────────────────────────

    # 用於偵測已有中文數字編號的正規式，例如 "一、" "二、"
    _RE_CN_NUM = re.compile(r"^[一二三四五六七八九十]{1,3}、")
    # 用於偵測已有次層級編號，例如 "(一)" "（一）"
    _RE_CN_SUB = re.compile(r"^[（(][一二三四五六七八九十]{1,3}[）)]")
    # 用於偵測阿拉伯數字編號，例如 "1." "2."
    _RE_ARABIC = re.compile(r"^\d+[.、]")

    def _auto_number(self, lines: list[str]) -> list[str]:
        """自動編號：將多項說明轉換為「一、二、三」及多層級編號。

        規則：
        - 若行已有「一、」「(一)」「1.」等編號，保持原樣
        - 若段落有多行實質內容（≥2 行），且都沒有編號，自動加上「一、二、三」
        - 子層級（以空格或 Tab 縮排的行）自動加上「(一)(二)(三)」
        - 三級（雙重縮排）自動加上「1. 2. 3.」
        """
        # 先清理得到非空行列表
        cleaned = []
        for line in lines:
            stripped = line.rstrip()
            if stripped:
                cleaned.append(stripped)

        if len(cleaned) < 2:
            return lines  # 單行不需要編號

        # 檢查是否已經有編號
        has_existing_numbering = any(
            self._RE_CN_NUM.match(self._clean_line(l))
            or self._RE_CN_SUB.match(self._clean_line(l))
            or self._RE_ARABIC.match(self._clean_line(l))
            for l in cleaned
        )
        if has_existing_numbering:
            return lines  # 已有編號，不覆蓋

        # 分析縮排層級
        result = []
        level1_idx = 0  # 一級計數
        level2_idx = 0  # 二級計數
        level3_idx = 0  # 三級計數

        for line in lines:
            stripped = line.rstrip()
            if not stripped:
                result.append(line)
                continue

            # 判斷縮排深度
            leading = len(line) - len(line.lstrip())
            # Tab 視為 4 個空格
            effective_indent = leading + line[:leading].count('\t') * 3

            if effective_indent >= 8:
                # 三級：1. 2. 3.
                level3_idx += 1
                result.append(f"{level3_idx}. {stripped.lstrip()}")
            elif effective_indent >= 2:
                # 二級：(一)(二)(三)
                if level2_idx < len(CHINESE_NUMBERS):
                    prefix = f"（{CHINESE_NUMBERS[level2_idx]}）"
                else:
                    prefix = f"（{level2_idx + 1}）"
                level2_idx += 1
                level3_idx = 0  # 重設三級計數
                result.append(f"{prefix}{stripped.lstrip()}")
            else:
                # 一級：一、二、三
                if level1_idx < len(CHINESE_NUMBERS):
                    prefix = f"{CHINESE_NUMBERS[level1_idx]}、"
                else:
                    prefix = f"{level1_idx + 1}、"
                level1_idx += 1
                level2_idx = 0  # 重設二級計數
                level3_idx = 0
                result.append(f"{prefix}{stripped}")

        return result

    def _write_attachments(self, doc: Document, sections: dict[str, str]) -> None:
        """寫入附件和參考來源。"""
        attachment_keys = [
            ("attachments", "附件："),
            ("references", "參考來源："),
        ]

        for key, label in attachment_keys:
            content = sections.get(key)
            if not content:
                continue

            doc.add_paragraph()
            p_att = doc.add_paragraph()
            self._set_paragraph_spacing(p_att)
            run_att = p_att.add_run(label)
            self._set_font(run_att, self._font_title, self._size_body, bold=True)

            # 逐行處理，與 _write_body 一致，避免多行內容中
            # 後續行的 markdown 標記（如 "- "）無法被清除
            for line in content.split("\n"):
                clean_content = self._clean_line(line)
                if clean_content:
                    p_content = doc.add_paragraph(clean_content)
                    self._set_paragraph_spacing(p_content)
                    for run in p_content.runs:
                        self._set_font(run, self._font_body, self._size_meta)

    def _write_qa_report(self, doc: Document, qa_report: str) -> None:
        """寫入 QA 報告附件頁。"""
        doc.add_page_break()
        p_qa_title = doc.add_paragraph()
        self._set_paragraph_spacing(p_qa_title)
        run_qa = p_qa_title.add_run("附件：AI 品質保證報告 (QA Report)")
        self._set_font(run_qa, self._font_title, self._size_section, bold=True)

        # 清理特殊字元和 Markdown 標記
        clean_qa = self._sanitize_text(qa_report)
        clean_qa = clean_markdown_artifacts(clean_qa)
        p_log = doc.add_paragraph(clean_qa)
        for run in p_log.runs:
            run.font.name = FONT_LOG
            run.font.size = Pt(FONT_SIZE_LOG)
